"""
Enterprise IAM Textbook — Comprehensive Word Document Generator
Reads all scripts from disk, embeds full code, prose, labs, ADRs, and case study.
Output: IAM_Textbook_Enterprise_Architecture.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent
OUTPUT_FILE = BASE / "IAM_Textbook_Enterprise_Architecture.docx"

REF_LIST = [
    "Microsoft. Active Directory Domain Services Overview. Microsoft Learn, 2024.",
    "Samba Team. Samba FSMO Roles Reference. samba.org, 2023.",
    "Royce, J. Kerberos: The Network Authentication Protocol. MIT, RFC 4120, 2005.",
    "Metcalf, S. Attacking Active Directory. adsecurity.org, 2023.",
    "Microsoft. Certificate Services Architecture. Microsoft Learn, 2024.",
    "Baines, O. Certified Pre-Owned: Abusing AD CS. SpecterOps, 2021.",
    "NIST. Digital Identity Guidelines. SP 800-63-3, 2017.",
    "Gartner. Magic Quadrant for Privileged Access Management. 2024.",
    "CyberArk. Privileged Access Management Reference Architecture. cyberark.com, 2024.",
    "HashiCorp. Vault Architecture Guide. developer.hashicorp.com, 2024.",
    "NIST. Zero Trust Architecture. SP 800-207, 2020.",
    "Microsoft. Conditional Access Overview. Microsoft Learn, 2024.",
    "Kindervag, J. No More Chewy Centers: Zero Trust. Forrester, 2010.",
    "SailPoint. Identity Governance Reference Architecture. sailpoint.com, 2024.",
    "ISACA. Segregation of Duties in IT. ISACA Journal, 2023.",
    "IETF. SCIM: System for Cross-domain Identity Management. RFC 7644, 2015.",
    "Microsoft. Microsoft Entra ID Documentation. Microsoft Learn, 2024.",
    "Microsoft. Microsoft Graph API Reference. Microsoft Learn, 2024.",
    "Auth0. Auth0 Documentation. auth0.com, 2024.",
    "Okta. Customer Identity and Access Management. okta.com, 2024.",
    "Microsoft. Microsoft Defender XDR Documentation. Microsoft Learn, 2024.",
    "Microsoft. Microsoft Sentinel KQL Reference. Microsoft Learn, 2024.",
    "MITRE. ATT&CK Framework. attack.mitre.org, 2024.",
    "NIST. AI Risk Management Framework. NIST AI 100-1, 2023.",
    "OWASP. LLM Top 10 for Large Language Model Applications. owasp.org, 2024.",
    "Anthropic. Claude API Documentation. anthropic.com, 2024.",
    "Gartner. AI TRiSM Framework. Gartner, 2023.",
    "ISO/IEC 42001. AI Management System Standard. ISO, 2023.",
]


def read_script(rel_path: str) -> str:
    p = BASE / rel_path
    if p.exists():
        return p.read_text(encoding="utf-8", errors="replace")
    return f"# File not found: {rel_path}\n"


def set_font(run, name="Segoe UI", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_toc(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("TABLE OF CONTENTS")
    set_font(run, bold=True, size=13)
    doc.add_paragraph()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    toc_para = doc.add_paragraph()
    run_elem = toc_para.add_run()
    run_elem._r.append(fldChar_begin)
    run_elem._r.append(instrText)
    run_elem._r.append(fldChar_separate)
    run_elem._r.append(fldChar_end)


def shade_cell(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_code_block(doc, code: str, filename: str = ""):
    if filename:
        lbl = doc.add_paragraph()
        r = lbl.add_run(filename)
        set_font(r, name="Segoe UI", size=9, bold=True, color=(80, 80, 80))
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F2F2F2")
    cell.paragraphs[0].clear()
    lines = code.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_font(r, name="Courier New", size=8)
    doc.add_paragraph()


def add_h1(doc, text: str):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(0, 70, 127))
    return p


def add_h2(doc, text: str):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=(31, 73, 125))
    return p


def add_h3(doc, text: str):
    p = doc.add_paragraph()
    p.style = "Heading 3"
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(54, 96, 146))
    return p


def add_body(doc, text: str, cite_nums: list = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r)
    if cite_nums:
        for n in cite_nums:
            rc = p.add_run(f" [{n}]")
            set_font(rc, size=8, color=(0, 0, 200))
    return p


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    r = p.add_run(text)
    set_font(r)


def analogy_box(doc, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF2CC")
    cell.paragraphs[0].clear()
    tp = cell.paragraphs[0]
    rt = tp.add_run(f"Analogy -- {title}")
    set_font(rt, size=10, bold=True, color=(127, 96, 0))
    bp = cell.add_paragraph()
    rb = bp.add_run(text)
    set_font(rb, size=10, italic=True)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────
def build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ENTERPRISE IDENTITY AND ACCESS MANAGEMENT")
    set_font(r, size=20, bold=True, color=(0, 70, 127))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("A Practitioner's Reference Textbook")
    set_font(r2, size=14, italic=True, color=(54, 96, 146))

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(
        "From Active Directory Foundations to AI-Driven Security Operations\n"
        "Including Complete Script Library, Lab Procedures, and Architecture Decision Records"
    )
    set_font(rs, size=11, color=(80, 80, 80))

    doc.add_paragraph()
    doc.add_paragraph()
    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = auth.add_run("Cybersec-by-Laurel  |  Enterprise IAM Portfolio Series")
    set_font(ra, size=10, bold=True)

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = ver.add_run("Edition 1.0  |  2026")
    set_font(rv, size=10, color=(100, 100, 100))
    add_page_break(doc)


# ─────────────────────────────────────────────
# PREFACE
# ─────────────────────────────────────────────
def build_preface(doc):
    add_h1(doc, "Preface")
    add_body(doc,
        "This textbook grew out of a practical need: enterprise IAM practitioners lack a single "
        "reference that spans the full stack -- from on-premises Active Directory forest design to "
        "AI-assisted access reviews. Most resources cover one layer in depth while ignoring the "
        "integration seams where real-world incidents occur. This work aims to fill that gap.")
    add_body(doc,
        "Vertex Health Systems -- an 85,000-employee, 28-hospital integrated health network -- "
        "serves as the running case study. Every architectural decision, every script, and every "
        "lab exercise is grounded in scenarios that practitioners at organisations of this scale "
        "encounter regularly.")
    add_body(doc,
        "Each chapter follows a deliberate structure: analogy first (to anchor the concept), "
        "theory second, implementation third, and attack-surface analysis fourth. Scripts are "
        "presented in full -- not as excerpts -- because understanding the complete implementation "
        "is the only way to reason about its security properties.")
    add_body(doc,
        "The 48 lab exercises are designed for a Windows Server 2022 lab environment with "
        "Microsoft Entra ID P2 licensing. Most PowerShell scripts run read-only by default; "
        "destructive operations require explicit parameter flags. Python scripts require "
        "Python 3.11+ and the anthropic SDK for Chapter 12 AI exercises.")
    add_body(doc,
        "Inline citations appear as blue bracketed numbers referencing the full bibliography "
        "at the end of the document. Architecture Decision Records (ADRs) document the reasoning "
        "behind major technology choices so that future teams can revisit them with full context.")
    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 1
# ─────────────────────────────────────────────
def chapter_01(doc):
    add_h1(doc, "Chapter 1 -- Identity Foundations and the IAM Discipline")

    analogy_box(doc,
        "The Hotel Key Card",
        "Authentication is proving you are a registered guest (identity). Authorisation is the "
        "set of floors and rooms your key card unlocks (entitlements). The front-desk system that "
        "records every door touch is auditing. IAM is the design of the entire key-card ecosystem: "
        "who issues cards, how long they last, how lost cards are revoked, and how the audit log "
        "is reviewed.")

    add_h2(doc, "1.1 What is IAM?")
    add_body(doc,
        "Identity and Access Management (IAM) is the discipline of ensuring that the right "
        "principals have the right access to the right resources at the right time -- and that "
        "every access event is logged for subsequent review.", [7])
    add_body(doc,
        "IAM spans four functional domains: Authentication (AuthN) -- verifying who you are; "
        "Authorisation (AuthZ) -- determining what you may do; Administration -- provisioning, "
        "modifying, and deprovisioning identities and entitlements; and Auditing -- recording "
        "and reviewing all identity-related events.")

    add_h2(doc, "1.2 The IAM Stack")
    for item in [
        "Directory Services -- Active Directory, Entra ID, LDAP",
        "PKI / Credential Infrastructure -- certificates, smart cards, FIDO2",
        "Privileged Access Management (PAM) -- just-in-time elevation, session recording",
        "Zero Trust Network Access -- conditional access, continuous evaluation",
        "Identity Governance and Administration (IGA) -- joiner/mover/leaver, access reviews",
        "Federation and CIAM -- SAML, OIDC, social login, Auth0",
        "Threat Detection -- SIEM, SOAR, KQL analytics, AI-assisted review",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "1.3 Compliance Drivers")
    add_body(doc,
        "IAM controls satisfy the access-control requirements of every major compliance "
        "framework: SOC 2 CC6.x, HIPAA Section 164.312(a), PCI DSS Requirement 7-8, NIST SP 800-53 "
        "AC family, and ISO 27001 Annex A.9. Understanding which IAM control satisfies which "
        "requirement prevents both over-engineering and gaps.", [7])

    add_h2(doc, "1.4 The Vertex Health Systems Case Study")
    add_body(doc,
        "Vertex Health Systems operates 28 hospitals across four US states with 85,000 employees, "
        "12,000 contractor identities, and 4.2 million patient portal accounts. The organisation "
        "runs a legacy on-premises Active Directory forest (vertexhealth.local) with 6 child "
        "domains inherited from acquisition activity. The IAM programme documented in this "
        "textbook consolidates this estate into a modern hybrid identity architecture.")
    add_body(doc,
        "Business requirements driving the programme: (1) eliminate standing privileged access "
        "by 2026; (2) achieve phishing-resistant MFA for all clinical staff; (3) reduce "
        "access-review cycle time from 90 days to 14 days; (4) pass HIPAA security-rule audit "
        "with zero access-control findings.")

    add_h2(doc, "1.5 How to Use This Textbook")
    add_body(doc,
        "Chapters 2-4 address the on-premises identity layer: Active Directory, PKI, and "
        "authentication protocols. Chapters 5-6 cover privileged access and Zero Trust. "
        "Chapters 7-9 address governance, Entra ID, and customer identity. Chapter 10 covers "
        "threat detection. Chapter 11 is the CAD capstone design. Chapter 12 introduces "
        "AI-assisted IAM and the emerging threat landscape of AI-enabled attacks.")
    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 2
# ─────────────────────────────────────────────
def chapter_02(doc):
    add_h1(doc, "Chapter 2 -- Active Directory and Directory Services")

    analogy_box(doc,
        "The Library Card Catalogue",
        "Active Directory is a library where every book (resource) has a catalogue card (object "
        "attribute), every borrower (user) has a library card (account), and the librarian "
        "(domain controller) maintains the authoritative index. FSMO roles are specialist "
        "librarians: one controls new card issuance (RID Master), one ensures the index is "
        "never corrupted by simultaneous edits (PDC Emulator), and one manages the master "
        "catalogue across branches (Schema Master).")

    add_h2(doc, "2.1 Active Directory Architecture")
    add_body(doc,
        "Active Directory Domain Services (AD DS) is a hierarchical, distributed database "
        "storing identity objects -- users, computers, groups, and service accounts -- and "
        "their attributes. It is organised into forests (security boundaries), domains "
        "(replication and policy units), and Organisational Units (OU) for delegation.", [1])
    add_body(doc,
        "A forest is the outermost security boundary. All domains within a forest share "
        "a common schema and global catalogue. Trust relationships allow authentication "
        "across domain boundaries; SID filtering controls which security identifiers are "
        "honoured across those trusts.", [1])

    add_h2(doc, "2.2 FSMO Roles")
    add_body(doc,
        "Flexible Single Master Operations (FSMO) roles prevent write conflicts in a "
        "multi-master replication environment. Five roles exist: Schema Master and Domain "
        "Naming Master are forest-wide (one each); PDC Emulator, RID Master, and "
        "Infrastructure Master are per-domain.", [1, 2])
    for role, desc in [
        ("Schema Master", "Controls schema modifications. One per forest."),
        ("Domain Naming Master", "Authorises adding/removing domains. One per forest."),
        ("PDC Emulator", "Time source, password lockout authority, legacy authentication fallback."),
        ("RID Master", "Allocates RID pools to DCs for SID construction."),
        ("Infrastructure Master", "Updates cross-domain group membership references."),
    ]:
        add_bullet(doc, f"{role}: {desc}")

    add_h2(doc, "2.3 Kerberos Authentication")
    add_body(doc,
        "Kerberos v5 is the default authentication protocol in Active Directory environments. "
        "The Key Distribution Centre (KDC), running on every domain controller, issues "
        "Ticket Granting Tickets (TGTs) on initial authentication and Service Tickets on "
        "resource access requests.", [3])
    analogy_box(doc,
        "The Theme-Park Wristband",
        "Your initial login produces a TGT -- like the wristband you receive at the theme-park "
        "entrance. When you want to ride a specific attraction (access a resource), you show "
        "your wristband to get a ride ticket (Service Ticket). The ride operator never needs "
        "to check your ID again -- only your ride ticket.")

    add_h2(doc, "2.4 Attack Surface: Kerberos and Delegation")
    add_body(doc,
        "Three Kerberos attack classes dominate enterprise incident reports. Kerberoasting "
        "extracts service tickets for service accounts (SPNs) and cracks them offline -- "
        "effective against weak passwords. AS-REP Roasting targets accounts with "
        "pre-authentication disabled. Unconstrained delegation allows any service to "
        "impersonate any user to any resource -- a critical misconfiguration.", [4])

    add_h2(doc, "2.5 Group Policy and OU Design")
    add_body(doc,
        "Group Policy Objects (GPOs) apply security configuration -- password policy, "
        "software restriction, audit settings -- to OUs. GPO inheritance can be blocked "
        "or enforced, creating audit complexity. Unlinked GPOs accumulate over time and "
        "represent dead configuration that complicates troubleshooting.", [1])

    add_h2(doc, "2.6 Script Library -- Chapter 2")
    add_body(doc,
        "The following nine scripts implement the Active Directory audit and hardening "
        "procedures described in this chapter. Each script is designed to run read-only "
        "by default; modification operations require explicit parameters.")

    scripts_02 = [
        ("scripts/chapter-02-directory/Get-FSMORoleHolders.ps1",
         "Script 2-1: FSMO Role Holder Audit",
         "Enumerates all five FSMO role holders, checks co-location with Global Catalogue, "
         "and exports a CSV for change-management records.", [1, 2]),
        ("scripts/chapter-02-directory/Find-DangerousAccountFlags.ps1",
         "Script 2-2: Dangerous UAC Flag Detection",
         "Scans all user accounts for high-risk User Account Control bitmask flags including "
         "unconstrained delegation (TRUSTED_FOR_DELEGATION), password not required "
         "(PASSWD_NOTREQD), and DES-only encryption (USE_DES_KEY_ONLY).", [1, 4]),
        ("scripts/chapter-02-directory/Get-KerberosAttackSurface.ps1",
         "Script 2-3: Kerberos Attack Surface Enumeration",
         "Identifies Kerberoastable accounts (SPNs on user objects), AS-REP Roastable "
         "accounts, unconstrained delegation targets, and krbtgt password age.", [3, 4]),
        ("scripts/chapter-02-directory/Measure-GroupNestingDepth.ps1",
         "Script 2-4: Group Nesting Depth Analysis",
         "Recursively measures group nesting depth with cycle detection. Deeply nested "
         "groups (>4 levels) cause authentication token bloat and authorisation failures.", [1]),
        ("scripts/chapter-02-directory/Find-EmptyAndStaleGroups.ps1",
         "Script 2-5: Empty and Stale Group Detection",
         "Identifies empty groups and groups unchanged for more than 365 days -- primary "
         "candidates for cleanup in IGA remediation campaigns.", [1]),
        ("scripts/chapter-02-directory/Get-GPOCoverageAudit.ps1",
         "Script 2-6: GPO Coverage Audit",
         "Reports unlinked GPOs and OUs with Block Inheritance set. Both conditions "
         "indicate configuration drift requiring review.", [1]),
        ("scripts/chapter-02-directory/Get-NTLMUsageAudit.ps1",
         "Script 2-7: NTLM Usage Audit",
         "Checks LmCompatibilityLevel registry value and queries Event 8004 volume to "
         "identify NTLM usage that should be migrated to Kerberos.", [1, 3]),
        ("scripts/chapter-02-directory/Get-TrustSIDFilteringAudit.ps1",
         "Script 2-8: Trust SID Filtering Audit",
         "Enumerates all domain and forest trusts, verifies SID filtering status, "
         "and flags trusts with filtering disabled (privilege escalation risk).", [1, 4]),
        ("scripts/chapter-02-directory/Get-ADHealthDashboard.ps1",
         "Script 2-9: Active Directory Health Dashboard",
         "Comprehensive health check covering replication status, privileged group membership, "
         "FSMO holders, krbtgt password age, and DC operating system versions.", [1, 2, 3]),
    ]

    for path, title, desc, refs in scripts_02:
        add_h3(doc, title)
        add_body(doc, desc, refs)
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h2(doc, "2.7 Lab Exercises -- Chapter 2")
    labs_02 = [
        ("Lab 2-A", "FSMO Role Transfer",
         "1. Open Active Directory Users and Computers as Domain Admin.\n"
         "2. Right-click domain root > Operations Masters.\n"
         "3. Document current RID Master holder.\n"
         "4. Run Get-FSMORoleHolders.ps1 and verify output matches GUI.\n"
         "5. Using Active Directory Module: Move-ADDirectoryServerOperationMasterRole "
         "-Identity DC02 -OperationMasterRole RIDMaster.\n"
         "6. Re-run script to confirm transfer.\n"
         "7. Transfer back to original holder."),
        ("Lab 2-B", "Kerberoast Simulation",
         "1. Create a service account: New-ADUser -Name svc_sql -Enabled $true.\n"
         "2. Set an SPN: setspn -A MSSQLSvc/sql01.corp.local:1433 svc_sql.\n"
         "3. Request a service ticket via .NET IdentityModel.\n"
         "4. Run Get-KerberosAttackSurface.ps1 -- verify svc_sql appears.\n"
         "5. Remediation: set a 127-character random password on svc_sql; "
         "consider gMSA replacement."),
        ("Lab 2-C", "GPO Inheritance Audit",
         "1. Run Get-GPOCoverageAudit.ps1.\n"
         "2. Identify any OUs with Block Inheritance.\n"
         "3. Review which GPOs are unlinked.\n"
         "4. For each unlinked GPO, determine if it can be deleted or should be linked.\n"
         "5. Document findings in the Vertex Health change log."),
        ("Lab 2-D", "NTLM Deprecation Baseline",
         "1. Run Get-NTLMUsageAudit.ps1 on each domain controller.\n"
         "2. Review Event 8004 count (NTLM pass-through authentications).\n"
         "3. Identify top NTLM-using workstations from Security log.\n"
         "4. Map NTLM usage to application owners.\n"
         "5. Create a migration plan to Kerberos or modern auth for each application."),
    ]
    for lab_id, lab_title, procedure in labs_02:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 3
# ─────────────────────────────────────────────
def chapter_03(doc):
    add_h1(doc, "Chapter 3 -- Public Key Infrastructure and Certificate Services")

    analogy_box(doc,
        "The Notary Public",
        "A Certificate Authority is a trusted notary. When a notary stamps a document, "
        "everyone who trusts that notary also trusts the document. The Root CA is the "
        "Grand Notary whose stamp is pre-installed in every browser and operating system. "
        "Intermediate CAs are regional offices -- they can stamp documents because the "
        "Grand Notary has stamped their authority certificate.")

    add_h2(doc, "3.1 PKI Hierarchy")
    add_body(doc,
        "A PKI hierarchy distributes trust. An offline Root CA issues only Intermediate CA "
        "certificates and is kept powered-off in a physically secure location -- this protects "
        "the root key even if online systems are compromised. Issuing CAs are online and "
        "issue end-entity certificates (users, computers, servers).", [5, 7])
    add_body(doc,
        "Vertex Health PKI decision: two-tier hierarchy. One offline Root CA (kept in a "
        "bank vault with HSM). Two online Issuing CAs for redundancy -- one per primary "
        "data centre. Certificate lifetime: Root 20 years, Issuing CA 10 years, "
        "end-entity 1 year (see ADR-003).")

    add_h2(doc, "3.2 Active Directory Certificate Services (AD CS)")
    add_body(doc,
        "AD CS integrates PKI with Active Directory, enabling auto-enrolment for computers "
        "and users. Certificate templates define what can be requested, who can request it, "
        "and what the certificate can be used for. Misconfigured templates are the root "
        "cause of ESC-class vulnerabilities.", [5, 6])

    add_h2(doc, "3.3 ESC Vulnerabilities")
    add_body(doc,
        "Certified Pre-Owned (2021) documented eight classes of AD CS misconfigurations "
        "(ESC1-ESC8) allowing privilege escalation or authentication bypass. ESC1: enrolee "
        "can specify Subject Alternative Name (SAN) -- allows impersonating any user. "
        "ESC2: template allows any purpose and has enrolee-supplied subject. "
        "ESC8: NTLM relay to HTTP-exposed certificate enrolment endpoint.", [6])
    analogy_box(doc,
        "The Blank Cheque",
        "ESC1 is equivalent to being handed a blank cheque and being told to fill in your "
        "own name as the payee. The certificate template lets the requester specify whose "
        "identity the certificate will assert -- so any user can claim to be the Domain Admin.")

    add_h2(doc, "3.4 Certificate Lifecycle Management")
    add_body(doc,
        "Certificates expire. Auto-enrolment reduces manual renewal burden for domain-joined "
        "systems. Non-domain-joined systems (IoT, third-party apps) require ACME protocol "
        "or manual renewal processes. Certificate revocation via CRL and OCSP must be "
        "accessible from all client networks -- a common gap in segmented environments.", [5])

    add_h2(doc, "3.5 Script Library -- Chapter 3")

    scripts_03 = [
        ("scripts/chapter-03-pki/Find-VulnerableCertTemplates.ps1",
         "Script 3-1: Vulnerable Certificate Template Detection (ESC1/ESC2/ESC3)",
         "Enumerates all certificate templates in AD, checks enrollment permissions, "
         "SAN supply flags, and EKU settings to identify ESC1/2/3-class vulnerabilities.", [5, 6]),
        ("scripts/chapter-03-pki/Get-CAInfrastructureAudit.ps1",
         "Script 3-2: CA Infrastructure Audit",
         "Audits all CAs in the forest: certificate expiry, HTTP certsrv exposure, "
         "and web enrolment endpoint availability (ESC8 precondition).", [5, 6]),
        ("scripts/chapter-03-pki/Get-CertificateExpiryReport.ps1",
         "Script 3-3: Certificate Expiry Monitoring",
         "Uses certutil to enumerate issued certificates and reports those expiring "
         "within 60 days (warning) or 30 days (critical).", [5]),
        ("scripts/chapter-03-pki/Test-ESC8Vulnerability.ps1",
         "Script 3-4: ESC8 Vulnerability Test",
         "Tests whether the CA web enrolment endpoint (certsrv) is accessible over HTTP "
         "and returns an NTLM challenge -- the two preconditions for ESC8 relay attacks.", [6]),
    ]

    for path, title, desc, refs in scripts_03:
        add_h3(doc, title)
        add_body(doc, desc, refs)
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h3(doc, "OpenSSL Reference -- PKI Operations")
    add_body(doc,
        "The following reference covers the 15 most common OpenSSL certificate operations "
        "performed by IAM engineers: key generation, CSR creation, chain validation, "
        "OCSP checking, PKCS12 bundling, CRL inspection, and TLS testing.")
    openssl_md = read_script("scripts/chapter-03-pki/openssl-reference.md")
    add_code_block(doc, openssl_md, "openssl-reference.md")

    add_h2(doc, "3.6 Lab Exercises -- Chapter 3")
    labs_03 = [
        ("Lab 3-A", "ESC1 Discovery and Remediation",
         "1. Run Find-VulnerableCertTemplates.ps1.\n"
         "2. Review output for templates with ENROLLEE_SUPPLIES_SUBJECT flag.\n"
         "3. For each vulnerable template: open Certificate Templates Console, "
         "locate template, check Request Handling > Supply in the request.\n"
         "4. Remediation: uncheck Supply in the request; require approval for Subject.\n"
         "5. Document templates remediated in Vertex Health PKI remediation log."),
        ("Lab 3-B", "Certificate Expiry Dashboard",
         "1. Run Get-CertificateExpiryReport.ps1 against both Issuing CAs.\n"
         "2. Identify certificates expiring within 30 days.\n"
         "3. Map each certificate to its owning system and application team.\n"
         "4. Initiate renewal requests for critical certificates.\n"
         "5. Validate renewal via: certutil -verify <cert.cer>."),
        ("Lab 3-C", "ESC8 Remediation",
         "1. Run Test-ESC8Vulnerability.ps1.\n"
         "2. If certsrv responds on HTTP: add HTTPS binding in IIS Manager, "
         "redirect HTTP to HTTPS.\n"
         "3. Disable NTLM on certsrv: require extended protection for authentication.\n"
         "4. Re-run Test-ESC8Vulnerability.ps1 to confirm remediation.\n"
         "5. Consider disabling web enrolment entirely if ADCS auto-enrolment is sufficient."),
    ]
    for lab_id, lab_title, procedure in labs_03:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 4
# ─────────────────────────────────────────────
def chapter_04(doc):
    add_h1(doc, "Chapter 4 -- Authentication Protocols: SAML, OIDC, and Federation")

    analogy_box(doc,
        "The Diplomatic Passport",
        "Federation is like international diplomatic recognition. Your home country (Identity "
        "Provider) issues your passport (SAML assertion or ID token). When you arrive at a "
        "foreign country (Service Provider), they trust your passport because your country and "
        "theirs have a bilateral agreement (federation trust). You never show your password -- "
        "only your passport.")

    add_h2(doc, "4.1 SAML 2.0")
    add_body(doc,
        "Security Assertion Markup Language (SAML) 2.0 is an XML-based federation protocol "
        "widely used in enterprise SSO. The Identity Provider (IdP) authenticates the user "
        "and issues a signed XML assertion. The Service Provider (SP) validates the "
        "assertion signature using the IdP's public certificate.")
    add_body(doc,
        "SAML flows: SP-initiated (user starts at SP, gets redirected to IdP) and "
        "IdP-initiated (user starts at IdP portal). SP-initiated is preferred -- "
        "IdP-initiated assertions lack a nonce and are vulnerable to replay attacks.")

    add_h2(doc, "4.2 OpenID Connect and OAuth 2.0")
    add_body(doc,
        "OIDC layers an identity layer on top of OAuth 2.0. The Authorization Server issues "
        "an ID Token (JWT) asserting who the user is, and an Access Token authorising API "
        "calls. OAuth 2.0 alone handles authorisation; OIDC adds authentication semantics.", [19])
    add_body(doc,
        "Token types: ID Token (who you are, for the client), Access Token (what you can do, "
        "for the API), Refresh Token (obtain new access tokens without re-authentication). "
        "Always validate: issuer, audience, expiry, and signature.")

    add_h2(doc, "4.3 Microsoft AD FS vs Pass-Through Sync")
    add_body(doc,
        "AD FS federates on-premises identity to cloud services by issuing SAML/OIDC tokens "
        "from on-premises infrastructure. Password Hash Sync (PHS) replicates a hash of the "
        "password hash to Entra ID, enabling cloud authentication without on-premises "
        "dependency. PHS is preferred for resilience -- if on-premises fails, authentication "
        "continues.", [7])
    add_body(doc,
        "Vertex Health ADR-002 selected PHS over AD FS to eliminate the AD FS infrastructure "
        "dependency. Seamless SSO provides transparent Kerberos-based SSO from domain-joined "
        "machines without routing authentication on-premises.")

    add_h2(doc, "4.4 Phishing-Resistant Authentication")
    add_body(doc,
        "Password + SMS OTP is vulnerable to SIM-swap and real-time phishing proxies. "
        "Phishing-resistant authentication requires a cryptographic binding between the "
        "authentication ceremony and the legitimate origin: FIDO2/WebAuthn passkeys, "
        "certificate-based authentication (CBA), or Windows Hello for Business.", [7])

    add_h2(doc, "4.5 Lab Exercises -- Chapter 4")
    for lab_id, lab_title, procedure in [
        ("Lab 4-A", "SAML Assertion Inspection",
         "1. Configure a SAML application in Entra ID.\n"
         "2. Use SAML tracer browser extension to capture the assertion.\n"
         "3. Base64-decode and review: Issuer, Subject, Conditions, AttributeStatement, Signature.\n"
         "4. Verify the NotBefore and NotOnOrAfter timestamps.\n"
         "5. Confirm the certificate thumbprint matches the IdP signing certificate."),
        ("Lab 4-B", "OIDC Token Inspection",
         "1. Configure an OIDC app registration in Entra ID.\n"
         "2. Obtain an ID Token using the authorization code flow.\n"
         "3. Decode the JWT at jwt.ms.\n"
         "4. Verify: iss, aud, exp, iat, sub, and custom claims.\n"
         "5. Review the token lifetime configuration in the app registration."),
        ("Lab 4-C", "FIDO2 Registration",
         "1. In Entra ID > Authentication Methods, enable FIDO2 security keys.\n"
         "2. Register a FIDO2 key for a test user.\n"
         "3. Sign in using the FIDO2 key.\n"
         "4. Review the sign-in log -- confirm authenticationMethodsUsed = FIDO2.\n"
         "5. Test: attempt to sign in from a phishing-proxy URL -- confirm the key refuses."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 5
# ─────────────────────────────────────────────
def chapter_05(doc):
    add_h1(doc, "Chapter 5 -- Privileged Access Management")

    analogy_box(doc,
        "The Safety Deposit Box",
        "Standing privileged access is like keeping the safety deposit box key on your "
        "keyring at all times -- convenient but a constant theft target. Just-in-time "
        "access is the bank's time-vault: you request access, a banker escorts you at the "
        "appointed time, and the vault locks again when your session ends. The valuable "
        "items are equally accessible but the attack window is milliseconds, not years.")

    add_h2(doc, "5.1 Why Standing Privilege is Dangerous")
    add_body(doc,
        "Standing privileged accounts -- accounts that are permanently members of Domain "
        "Admins, Global Admins, or equivalent -- are the primary target of every advanced "
        "persistent threat actor. They represent persistent, always-available privilege "
        "that can be harvested via credential theft, lateral movement, or insider abuse.", [8])

    add_h2(doc, "5.2 Just-In-Time Access Models")
    add_body(doc,
        "JIT access grants privilege for a specific time window in response to an approved "
        "request. Three implementation models: (1) Entra PIM -- activates role membership "
        "for 1-8 hours with approval and MFA. (2) CyberArk -- checks out credentials from "
        "vault and rotates after session. (3) AD JIT Groups -- time-bounded AD group "
        "membership via scheduled task auto-revocation (Script 5-2).", [8, 9])

    add_h2(doc, "5.3 PowerShell Security and Logging")
    add_body(doc,
        "PowerShell is the most powerful administrative tool in Windows environments and "
        "therefore the most abused by attackers. Script Block Logging (Event 4104) records "
        "deobfuscated script content; Module Logging (4103) records pipeline execution; "
        "Transcription saves full session transcripts; AMSI Integration enables real-time "
        "content inspection by antimalware.")

    add_h2(doc, "5.4 LAPS and Local Administrator Management")
    add_body(doc,
        "Local Administrator Password Solution (LAPS) generates unique, random local "
        "administrator passwords for each managed endpoint and stores them in AD (Legacy "
        "LAPS) or Entra ID (Windows LAPS). This prevents lateral movement via shared "
        "local admin credentials -- a technique used in over 60% of ransomware incidents.", [9])

    add_h2(doc, "5.5 Break-Glass Accounts")
    add_body(doc,
        "Break-glass (emergency access) accounts are cloud-only Global Admin accounts "
        "excluded from all Conditional Access policies. Requirements: FIDO2-only "
        "authentication, no MFA method that can be lost, stored credentials in physically "
        "secured location, sign-in alerts monitored 24/7.")

    add_h2(doc, "5.6 Script Library -- Chapter 5")

    scripts_05 = [
        ("scripts/chapter-05-pam/Get-PIMAudit.ps1",
         "Script 5-1: Entra PIM Audit",
         "Reports permanent active role assignments (should be zero for privileged roles), "
         "counts eligible Global Admin assignments, and reviews recent role activations.", [8]),
        ("scripts/chapter-05-pam/New-JITGroupMembership.ps1",
         "Script 5-2: Just-in-Time AD Group Membership",
         "Creates a time-bounded AD group membership by adding a user and scheduling "
         "an automatic removal task after the specified duration.", [9]),
        ("scripts/chapter-05-pam/Set-PowerShellLogging.ps1",
         "Script 5-3: PowerShell Logging Configuration",
         "Enables Script Block Logging (4104), Module Logging (4103), and Transcription "
         "via registry keys and Group Policy-compatible settings.", [8]),
        ("scripts/chapter-05-pam/Find-SuspiciousPSCommands.ps1",
         "Script 5-4: Suspicious PowerShell Command Detection",
         "Queries Event 4104 for 14 attack signatures including mimikatz, AMSI bypass, "
         "BloodHound, Invoke-Kerberoast, and credential dumping patterns.", [4, 8]),
        ("scripts/chapter-05-pam/Get-LAPSAudit.ps1",
         "Script 5-5: LAPS Coverage Audit",
         "Reports computers with Legacy LAPS, Windows LAPS, and no LAPS coverage. "
         "Also identifies computers with expired LAPS passwords.", [9]),
        ("scripts/chapter-05-pam/Watch-BreakGlassAccounts.ps1",
         "Script 5-6: Break-Glass Account Monitor",
         "Validates break-glass account configuration: cloud-only check, FIDO2-only "
         "authentication methods, CA policy exclusion, and last sign-in activity.", [8]),
        ("scripts/chapter-05-pam/Get-KerberosAttackSurface.ps1",
         "Script 5-7: PAM-Focused Kerberos Attack Surface",
         "Identifies service accounts with SPNs that are also members of privileged groups "
         "-- the highest-risk Kerberoasting targets.", [3, 4]),
    ]

    for path, title, desc, refs in scripts_05:
        add_h3(doc, title)
        add_body(doc, desc, refs)
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h2(doc, "5.7 Lab Exercises -- Chapter 5")
    for lab_id, lab_title, procedure in [
        ("Lab 5-A", "PIM Role Activation",
         "1. In Entra ID, make a test user eligible for Security Reader role via PIM.\n"
         "2. Sign in as the test user and activate the role.\n"
         "3. Observe the approval workflow and MFA requirement.\n"
         "4. Run Get-PIMAudit.ps1 -- verify no permanent active assignments.\n"
         "5. Let activation expire and verify the role is removed."),
        ("Lab 5-B", "JIT AD Group Membership",
         "1. Run New-JITGroupMembership.ps1 -UserSamAccountName testuser "
         "-GroupName 'Server-Admins' -DurationMinutes 30.\n"
         "2. Verify user is added to group: Get-ADGroupMember 'Server-Admins'.\n"
         "3. Verify scheduled task was created: Get-ScheduledTask -TaskName 'JIT-Remove-*'.\n"
         "4. Wait for task to run, verify user is removed.\n"
         "5. Review the JIT audit log entry."),
        ("Lab 5-C", "PowerShell Logging Verification",
         "1. Run Set-PowerShellLogging.ps1.\n"
         "2. Open a new PowerShell session.\n"
         "3. Execute: Get-ADUser -Filter * (a normal command).\n"
         "4. Open Event Viewer > Applications and Services Logs > PowerShell.\n"
         "5. Verify Event 4104 contains the deobfuscated script block.\n"
         "6. Run Find-SuspiciousPSCommands.ps1 to confirm detection pipeline works."),
        ("Lab 5-D", "Break-Glass Account Configuration",
         "1. Create a cloud-only account: breakglass@vertexhealth.onmicrosoft.com.\n"
         "2. Assign Global Administrator role permanently.\n"
         "3. Register a FIDO2 key -- the ONLY authentication method.\n"
         "4. Verify account is excluded from all Conditional Access policies.\n"
         "5. Run Watch-BreakGlassAccounts.ps1 to validate configuration.\n"
         "6. Store the recovery kit in the physical safe."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 6
# ─────────────────────────────────────────────
def chapter_06(doc):
    add_h1(doc, "Chapter 6 -- Zero Trust Architecture and Conditional Access")

    analogy_box(doc,
        "The Airport Security Checkpoint",
        "Traditional perimeter security is like a castle with a moat: get across the "
        "drawbridge (VPN) and you have free run of the interior. Zero Trust is like a "
        "modern airport: you pass through security to enter, but you also show your "
        "boarding pass (authorisation) at every gate. Each gate makes an independent "
        "decision -- it does not trust you just because you got past the entrance.")

    add_h2(doc, "6.1 Zero Trust Principles")
    add_body(doc,
        "Zero Trust, formalised by NIST SP 800-207, rests on three principles: "
        "(1) Never trust, always verify -- no implicit trust based on network location. "
        "(2) Assume breach -- design as if attackers are already inside. "
        "(3) Least privilege -- grant minimum access required for each operation.", [11])

    add_h2(doc, "6.2 Conditional Access Policy Design")
    add_body(doc,
        "Conditional Access (CA) is the Zero Trust policy engine in Entra ID. Policies "
        "evaluate signals -- user identity, device compliance, location, application, "
        "session risk -- and enforce controls: MFA, compliant device, approved app, "
        "or session restrictions.", [12])

    add_h2(doc, "6.3 Continuous Access Evaluation (CAE)")
    add_body(doc,
        "Traditional access tokens live for 1 hour regardless of risk changes. CAE "
        "enables near-real-time token revocation when the IdP sends a revocation event -- "
        "CAE-capable applications re-evaluate within seconds rather than waiting for "
        "token expiry. This closes the window between compromise detection and "
        "access termination.", [12])

    add_h2(doc, "6.4 Script Library -- Chapter 6")

    scripts_06 = [
        ("scripts/chapter-06-zerotrust/Get-CAPolicyGapAnalysis.ps1",
         "Script 6-1: Conditional Access Policy Gap Analysis",
         "Evaluates existing CA policies against the Vertex Health baseline tiers: "
         "Tier B (block legacy auth), Tier I (require MFA), Tier P (privileged role "
         "protection). Reports gaps for each tier.", [11, 12]),
        ("scripts/chapter-06-zerotrust/Measure-ZeroTrustMaturity.ps1",
         "Script 6-2: Zero Trust Maturity Assessment",
         "Interactive CISA 5-pillar Zero Trust Maturity Model self-assessment. "
         "Scores each pillar on a 1-4 scale and exports results to JSON.", [11, 13]),
        ("scripts/chapter-06-zerotrust/Test-CAEAndRevocation.ps1",
         "Script 6-3: CAE and Token Revocation Test",
         "Analyses sign-in logs for modern vs legacy authentication distribution "
         "and tests token revocation latency via Entra ID revocation APIs.", [12]),
    ]

    for path, title, desc, refs in scripts_06:
        add_h3(doc, title)
        add_body(doc, desc, refs)
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h2(doc, "6.5 Lab Exercises -- Chapter 6")
    for lab_id, lab_title, procedure in [
        ("Lab 6-A", "CA Policy Gap Analysis",
         "1. Run Get-CAPolicyGapAnalysis.ps1.\n"
         "2. Review the gap report for each tier (B/I/P).\n"
         "3. For each gap: create the missing CA policy in report-only mode first.\n"
         "4. Monitor for 7 days. Review impact analysis, then enable."),
        ("Lab 6-B", "Zero Trust Maturity Assessment",
         "1. Run Measure-ZeroTrustMaturity.ps1.\n"
         "2. Answer each pillar question honestly.\n"
         "3. Review the JSON output.\n"
         "4. Identify the two lowest-scoring pillars.\n"
         "5. Create a 90-day improvement plan for each pillar."),
        ("Lab 6-C", "Legacy Authentication Block",
         "1. In Entra ID Sign-in Logs, filter by Client App = Other clients.\n"
         "2. Identify users still using legacy auth protocols.\n"
         "3. Contact application owners for each identified application.\n"
         "4. Create CA policy blocking legacy auth in report-only mode.\n"
         "5. After 30-day monitoring, enable the block policy."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 7
# ─────────────────────────────────────────────
def chapter_07(doc):
    add_h1(doc, "Chapter 7 -- Identity Governance and Administration")

    analogy_box(doc,
        "The Employee Badge Audit",
        "IGA is the quarterly review where HR walks through the building and asks: "
        "does this person still work here? Do they still need access to this floor? "
        "Without IGA, access accumulates like clutter -- every project, every role change, "
        "every temporary assignment leaves a permanent access trace. IGA is the "
        "institutional memory that asks: why does this person still have this?")

    add_h2(doc, "7.1 The Joiner-Mover-Leaver Lifecycle")
    add_body(doc,
        "Every identity follows a lifecycle: Joiner (new employee -- provision access), "
        "Mover (role change -- modify access to match new role, remove access from old role), "
        "and Leaver (termination -- revoke all access within SLA). Access that accumulates "
        "from the Mover phase without removal is called access creep -- the primary "
        "driver of Segregation of Duties violations.")

    add_h2(doc, "7.2 Segregation of Duties")
    add_body(doc,
        "SoD prevents a single individual from having conflicting entitlements -- "
        "specifically, the ability to initiate and approve the same transaction. "
        "In financial systems, AP-Create-Vendor and AP-Approve-Payment conflict: "
        "one person holding both can create a fraudulent vendor and approve payment to it.", [14, 15])

    add_h2(doc, "7.3 Access Certifications")
    add_body(doc,
        "Access certification (access review) is the periodic process of having managers "
        "or application owners certify that entitlements are still appropriate. "
        "Certification frequency should match entitlement risk: quarterly for privileged "
        "access, annually for standard access. AI-assisted review (Chapter 12) accelerates "
        "this process by pre-populating recommendations.")

    add_h2(doc, "7.4 SCIM Provisioning")
    add_body(doc,
        "System for Cross-domain Identity Management (SCIM) 2.0 is the standard protocol "
        "for automated user provisioning between identity providers and applications. "
        "SCIM enables real-time provisioning on joiner events and deprovisioning on "
        "leaver events without manual intervention.", [16])
    analogy_box(doc,
        "The Automated Passport System",
        "SCIM is like an automated immigration system: when HR records a new hire, the "
        "system automatically notifies every downstream application -- without requiring a "
        "human to manually update each one. The leaver process triggers automatic "
        "deactivation everywhere simultaneously.")

    add_h2(doc, "7.5 Script Library -- Chapter 7")

    scripts_07 = [
        ("scripts/chapter-07-iga/Find-SoDViolations.ps1",
         "Script 7-1: Segregation of Duties Violation Detection",
         "CSV-driven SoD conflict detection. Reads a conflict matrix and entitlement "
         "assignments, identifies users with conflicting entitlement pairs, and calculates "
         "violation duration.", [14, 15]),
        ("scripts/chapter-07-iga/Invoke-LeaverProcessing.ps1",
         "Script 7-2: Automated Leaver Processing",
         "Implements the leaver workflow: WHO/WHAT/WHEN/WHY audit logging, 4-hour "
         "SLA enforcement, Entra ID session revocation, and account disable operations.", [14]),
        ("scripts/chapter-07-iga/Test-SCIMEndpoint.ps1",
         "Script 7-3: SCIM Endpoint Lifecycle Test",
         "Full SCIM 2.0 lifecycle test: Create user, Read user, Patch (update) user, "
         "Deactivate user (PATCH active=false), Delete user. Validates endpoint compliance.", [16]),
    ]

    for path, title, desc, refs in scripts_07:
        add_h3(doc, title)
        add_body(doc, desc, refs)
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h2(doc, "7.6 Lab Exercises -- Chapter 7")
    for lab_id, lab_title, procedure in [
        ("Lab 7-A", "SoD Violation Detection",
         "1. Create a SoD conflict matrix CSV (Zone1, Zone2, Severity).\n"
         "2. Create an entitlement assignment CSV (UserID, Entitlement, Zone).\n"
         "3. Run Find-SoDViolations.ps1 with your test data.\n"
         "4. Review output for HIGH and CRITICAL violations.\n"
         "5. For each violation: determine which entitlement to revoke based on role.\n"
         "6. Document remediation plan in the Vertex Health IGA remediation tracker."),
        ("Lab 7-B", "Leaver Processing Simulation",
         "1. Create a test user in AD and assign to several security groups.\n"
         "2. Run Invoke-LeaverProcessing.ps1 -UserSamAccountName testleaver.\n"
         "3. Verify: account disabled, group memberships removed, session revoked.\n"
         "4. Review the WHO/WHAT/WHEN/WHY audit log entry.\n"
         "5. Verify completion within the 4-hour SLA."),
        ("Lab 7-C", "SCIM Provisioning Test",
         "1. Configure a SCIM endpoint (use Postman or a test SCIM server).\n"
         "2. Run Test-SCIMEndpoint.ps1 against your endpoint.\n"
         "3. Verify all lifecycle operations: Create, Read, Patch, Deactivate, Delete.\n"
         "4. Review the SCIM server logs to confirm proper request format.\n"
         "5. Test error handling: attempt to create a duplicate user."),
        ("Lab 7-D", "Access Certification Campaign",
         "1. In Entra ID Identity Governance, create an Access Review for the "
         "Clinical-Apps-Users group.\n"
         "2. Set reviewer = Group Owner, duration = 14 days, frequency = Quarterly.\n"
         "3. As the group owner, review each member and make approve/deny decisions.\n"
         "4. Observe auto-apply: denied members are removed after campaign closes.\n"
         "5. Export results and calculate reviewer completion rate."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 8
# ─────────────────────────────────────────────
def chapter_08(doc):
    add_h1(doc, "Chapter 8 -- Microsoft Entra ID: Cloud Identity Platform")

    analogy_box(doc,
        "The Cloud Concierge",
        "Entra ID is the hotel concierge who knows every guest (identity), every room "
        "they are allowed to enter (authorisation), and every service they have used (audit). "
        "Unlike Active Directory (which is the on-premises building security desk), "
        "the Entra concierge operates across all cloud properties simultaneously -- "
        "regardless of whether the guest is in the building or calling from abroad.")

    add_h2(doc, "8.1 Entra ID vs Active Directory")
    add_body(doc,
        "Entra ID is not Active Directory in the cloud. AD is a Kerberos/LDAP directory "
        "designed for domain-joined devices on a LAN. Entra ID is a modern identity "
        "platform using OIDC/SAML/OAuth 2.0, designed for cloud applications and "
        "remote access. Entra Connect syncs identities from AD to Entra ID, enabling "
        "hybrid identity.", [17])

    add_h2(doc, "8.2 App Registrations and Service Principals")
    add_body(doc,
        "An App Registration defines an application's identity in Entra ID: credentials "
        "(secrets/certificates), permissions, and redirect URIs. A Service Principal is "
        "the instance of that application within a specific tenant. Application permissions "
        "(app roles) allow background processes to act without a user context; delegated "
        "permissions act on behalf of a signed-in user.", [17])

    add_h2(doc, "8.3 Guest Identity Management")
    add_body(doc,
        "Entra External Collaboration (B2B) allows partner and contractor identities "
        "from external tenants or social providers. Risks: guests with pending acceptance "
        "for >30 days (invitation phishing window), guests with no sign-in >90 days "
        "(stale identities), and guests without group membership (unused access).")

    add_h2(doc, "8.4 Microsoft Graph API")
    add_body(doc,
        "Microsoft Graph is the unified API for all Microsoft cloud services. Delta queries "
        "use a delta link to retrieve only changes since the last query -- essential for "
        "efficient synchronisation in large tenants.", [18])

    add_h2(doc, "8.5 Script Library -- Chapter 8")

    scripts_08 = [
        ("scripts/chapter-08-entra/Get-AppRegistrationSecurityAudit.ps1",
         "Script 8-1: App Registration Security Audit",
         "Identifies expiring client secrets (<30/60 day warning), apps without owners, "
         "multi-tenant applications, and apps with application (non-delegated) permissions.", [17]),
        ("scripts/chapter-08-entra/Get-StaleGuestAudit.ps1",
         "Script 8-2: Stale Guest Identity Audit",
         "Reports guests with pending invitation acceptance, no sign-in in >90 days, "
         "and guests with no group membership assignments.", [17]),
        ("scripts/chapter-08-entra/Get-GraphDeltaChanges.ps1",
         "Script 8-3: Microsoft Graph Delta Query",
         "Demonstrates delta link persistence for incremental identity synchronisation. "
         "First run performs full sync; subsequent runs retrieve only changes.", [18]),
    ]

    for path, title, desc, refs in scripts_08:
        add_h3(doc, title)
        add_body(doc, desc, refs)
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h2(doc, "8.6 Lab Exercises -- Chapter 8")
    for lab_id, lab_title, procedure in [
        ("Lab 8-A", "App Registration Security Audit",
         "1. Run Get-AppRegistrationSecurityAudit.ps1 against your tenant.\n"
         "2. Identify applications with secrets expiring within 30 days.\n"
         "3. Contact application owners for each expiring credential.\n"
         "4. Rotate the secret: New-AzADAppCredential -ObjectId <appId>.\n"
         "5. Review applications without owners -- assign owners or plan decommission."),
        ("Lab 8-B", "Guest Access Review",
         "1. Run Get-StaleGuestAudit.ps1.\n"
         "2. Review guests with pending invitations >30 days.\n"
         "3. Cancel stale invitations: Remove-MgUser -UserId <guestId>.\n"
         "4. Review guests with no sign-in >90 days with their sponsoring manager.\n"
         "5. Create a quarterly guest review process in Entra Identity Governance."),
        ("Lab 8-C", "Graph Delta Synchronisation",
         "1. Run Get-GraphDeltaChanges.ps1 -- observe full sync output.\n"
         "2. Create a new user in Entra ID.\n"
         "3. Re-run the script with the saved delta link -- observe incremental output.\n"
         "4. Verify only the new user appears in the delta response.\n"
         "5. Test with a user deletion -- verify the delta returns the deleted object."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 9
# ─────────────────────────────────────────────
def chapter_09(doc):
    add_h1(doc, "Chapter 9 -- Customer Identity and Access Management with Auth0")

    analogy_box(doc,
        "The Concert Venue Ticketing System",
        "CIAM is the public-facing ticketing system, not the employee badge system. "
        "Millions of concert-goers (customers) need to create accounts, buy tickets, "
        "and manage preferences -- all without IT department involvement. Auth0 handles "
        "social login (Log in with Google), multi-factor authentication, and organisation "
        "segregation (VIP section vs general admission).")

    add_h2(doc, "9.1 CIAM vs Workforce IAM")
    add_body(doc,
        "Customer IAM (CIAM) differs from workforce IAM in four dimensions: "
        "(1) Scale -- millions of users vs thousands of employees. "
        "(2) Self-service -- customers register themselves; employees are provisioned. "
        "(3) Friction tolerance -- customers abandon flows; employees comply with policy. "
        "(4) Identity verification -- customers need progressive profiling; employees "
        "are verified by HR.", [19, 20])

    add_h2(doc, "9.2 Auth0 Architecture")
    add_body(doc,
        "Auth0 is a cloud-native CIAM platform providing: Universal Login (hosted "
        "login UI), Actions (serverless event hooks), Custom Databases (migration), "
        "Organisations (B2B multi-tenancy), and Machine-to-Machine authentication. "
        "It implements OIDC/OAuth 2.0 and SAML.", [19])

    add_h2(doc, "9.3 Custom Database Migration")
    add_body(doc,
        "Legacy applications store passwords in custom databases, often with weak "
        "hashing (MD5, SHA1). Auth0 Custom Database with lazy migration enables "
        "zero-downtime migration: authenticate against the legacy database on first "
        "login, re-hash with bcrypt, store in Auth0. Subsequent logins use Auth0 directly.")
    analogy_box(doc,
        "The Lock Re-keying Migration",
        "Lazy migration is re-keying a building one tenant at a time. Each time a tenant "
        "uses their key (logs in), the locksmith (Auth0) swaps it for a modern key. "
        "Tenants who never return keep their old key -- until the migration deadline.")

    add_h2(doc, "9.4 Auth0 Actions")
    add_body(doc,
        "Auth0 Actions are serverless Node.js functions executing at specific points in "
        "the authentication pipeline. They replace the older Rules and Hooks. "
        "Use cases: claim enrichment, adaptive MFA, licence enforcement, migration triggers.")

    add_h2(doc, "9.5 Script Library -- Chapter 9")

    scripts_09 = [
        ("scripts/chapter-09-auth0/New-Auth0Organization.js",
         "Script 9-1: Create Auth0 Organisation",
         "Creates an Auth0 Organisation for a hospital tenant, enables the enterprise "
         "connection, and sends an administrator invitation.", [19]),
        ("scripts/chapter-09-auth0/Auth0TokenCache.js",
         "Script 9-2: Management API Token Cache",
         "In-memory access token cache with 60-second pre-expiry and single in-flight "
         "promise to prevent token stampede on concurrent requests.", [19]),
        ("scripts/chapter-09-auth0/custom-db-login.js",
         "Script 9-3: Custom Database Login with Lazy Migration",
         "Authenticates users against a legacy database. Detects MD5-salted vs bcrypt "
         "hashing and triggers migration to bcrypt on successful login.", [19]),
        ("scripts/chapter-09-auth0/action-post-login-enrich.js",
         "Script 9-4: Post-Login Action -- Claim Enrichment and Adaptive MFA",
         "Post-login Action that validates org membership, enriches claims with "
         "permissions, enforces adaptive MFA for high-risk operations, and flags "
         "legacy accounts for migration.", [19]),
        ("scripts/chapter-09-auth0/deploy-tenant.sh",
         "Script 9-5: Auth0 Tenant Deployment (Deploy CLI)",
         "Bash script wrapping the Auth0 Deploy CLI for export, import, and diff "
         "operations to manage tenant configuration as code.", [19]),
    ]

    for path, title, desc, refs in scripts_09:
        add_h3(doc, title)
        add_body(doc, desc, refs)
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h2(doc, "9.6 Lab Exercises -- Chapter 9")
    for lab_id, lab_title, procedure in [
        ("Lab 9-A", "Auth0 Organisation Setup",
         "1. Create an Auth0 tenant (free tier for lab).\n"
         "2. Enable Organisations feature.\n"
         "3. Run New-Auth0Organization.js with TENANT_DOMAIN, CLIENT_ID, CLIENT_SECRET.\n"
         "4. Verify the organisation appears in Auth0 dashboard.\n"
         "5. Log in as the invited admin and verify org membership."),
        ("Lab 9-B", "Custom Database Migration",
         "1. Create a test MySQL database with a users table.\n"
         "2. Insert test users with MD5-salted passwords.\n"
         "3. Configure Auth0 Custom Database connection.\n"
         "4. Implement the Login script using custom-db-login.js.\n"
         "5. Log in -- verify the user is migrated to Auth0 on first login.\n"
         "6. Log in again -- verify Auth0 handles authentication directly."),
        ("Lab 9-C", "Post-Login Action",
         "1. In Auth0, create a new Action (Flows > Login).\n"
         "2. Paste the content of action-post-login-enrich.js.\n"
         "3. Deploy the action and attach it to the Login flow.\n"
         "4. Log in as a test user.\n"
         "5. Inspect the ID token -- verify enriched claims appear.\n"
         "6. Test the legacy migration flag for users without bcrypt hashing."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 10
# ─────────────────────────────────────────────
def chapter_10(doc):
    add_h1(doc, "Chapter 10 -- Threat Detection: Microsoft Defender XDR and Sentinel")

    analogy_box(doc,
        "The Security Operations Centre as Air Traffic Control",
        "A SOC watching a SIEM is like air traffic control: hundreds of blips (events) "
        "on radar, each needing rapid classification. KQL analytics rules are the automated "
        "collision-avoidance system -- they flag patterns that human controllers must "
        "investigate immediately. SOAR playbooks are the auto-pilot: when the alarm sounds, "
        "first-response actions execute automatically while the human takes the helm.")

    add_h2(doc, "10.1 Microsoft Defender XDR")
    add_body(doc,
        "Microsoft Defender XDR integrates signals from Defender for Identity, Defender "
        "for Endpoint, Defender for Cloud Apps, and Defender for Office 365 into a unified "
        "incident view. Defender for Identity monitors AD and Entra ID for identity-based "
        "attack patterns.", [21])

    add_h2(doc, "10.2 Microsoft Sentinel and KQL")
    add_body(doc,
        "Microsoft Sentinel is a cloud-native SIEM/SOAR built on Azure Log Analytics. "
        "Kusto Query Language (KQL) uses a pipe-based syntax: each operator transforms "
        "the dataset, passing results to the next operator.", [22])

    add_h2(doc, "10.3 Identity Attack Patterns and KQL Detection")
    add_body(doc,
        "Seven KQL detection rules cover the primary identity attack patterns. "
        "Each rule is mapped to the MITRE ATT&CK framework and includes the specific "
        "event data it requires.", [22, 23])

    add_h2(doc, "10.4 Script Library -- Chapter 10")

    add_h3(doc, "Script 10-1: Zero Trust Baseline Measurement")
    add_body(doc,
        "Measures Zero Trust implementation status: MFA coverage percentage, "
        "phishing-resistant MFA percentage, risky users, legacy authentication volume, "
        "and CA policy count.", [11, 12])
    code = read_script("scripts/chapter-10-defender/Measure-ZeroTrustBaseline.ps1")
    add_code_block(doc, code, "Measure-ZeroTrustBaseline.ps1")

    kql_rules = [
        ("scripts/chapter-10-defender/sentinel-kql-rules/impossible-travel.kql",
         "KQL Rule 10-2: Impossible Travel Detection",
         "Detects sign-ins from geographically impossible locations using Haversine "
         "distance calculation. Flags sign-ins >500km apart within a 1-hour window."),
        ("scripts/chapter-10-defender/sentinel-kql-rules/global-admin-assignment.kql",
         "KQL Rule 10-3: Global Admin Role Assignment",
         "Detects assignment of 8 sensitive Entra ID roles including Global Administrator, "
         "Privileged Role Administrator, and Application Administrator. MITRE T1098."),
        ("scripts/chapter-10-defender/sentinel-kql-rules/kerberoasting-detection.kql",
         "KQL Rule 10-4: Kerberoasting Detection",
         "Detects Event 4769 with RC4 encryption type (0x17) at high volume -- "
         "more than 10 tickets in 15 minutes from a single source. MITRE T1558.003."),
        ("scripts/chapter-10-defender/sentinel-kql-rules/dcsync-detection.kql",
         "KQL Rule 10-5: DCSync Detection",
         "Detects Event 4662 with DS-Replication-Get-Changes access rights from "
         "non-domain-controller sources. MITRE T1003.006."),
        ("scripts/chapter-10-defender/sentinel-kql-rules/mass-user-modification.kql",
         "KQL Rule 10-6: Mass User Modification Detection",
         "Detects bulk user attribute modifications: more than 20 user changes in "
         "15 minutes by a single initiator. MITRE T1098."),
        ("scripts/chapter-10-defender/sentinel-kql-rules/service-principal-cred-add.kql",
         "KQL Rule 10-7: Service Principal Credential Addition",
         "Detects addition of client secrets or certificates to app registrations. "
         "MITRE T1098.001."),
        ("scripts/chapter-10-defender/sentinel-kql-rules/msol-anomaly-detection.kql",
         "KQL Rule 10-8: MSOL Anomaly Detection",
         "Detects unexpected operations by MSOL_ sync accounts and unexpected sign-in "
         "patterns -- both indicators of synchronisation account compromise."),
    ]

    for path, title, desc in kql_rules:
        add_h3(doc, title)
        add_body(doc, desc, [22, 23])
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h2(doc, "10.5 Lab Exercises -- Chapter 10")
    for lab_id, lab_title, procedure in [
        ("Lab 10-A", "KQL Threat Hunting",
         "1. Open Microsoft Sentinel > Logs.\n"
         "2. Paste the impossible-travel.kql query.\n"
         "3. Set time range to last 30 days.\n"
         "4. Review results -- identify any legitimate travel vs impossible cases.\n"
         "5. For each hit: review the user's travel schedule and IP geolocation.\n"
         "6. Create an Analytics Rule from the query to automate future detection."),
        ("Lab 10-B", "Zero Trust Baseline Measurement",
         "1. Run Measure-ZeroTrustBaseline.ps1.\n"
         "2. Review the JSON output.\n"
         "3. Identify the MFA coverage percentage.\n"
         "4. Target: >95% MFA, >50% phishing-resistant, 0 legacy auth.\n"
         "5. Create a 90-day roadmap to reach target metrics."),
        ("Lab 10-C", "Incident Response Simulation",
         "1. In Microsoft Sentinel, create a test incident from a KQL analytics rule.\n"
         "2. Assign the incident to yourself.\n"
         "3. Add a comment with initial triage findings.\n"
         "4. Run the SOAR playbook (if configured) to gather enrichment data.\n"
         "5. Close the incident with classification: True Positive / Benign Positive / "
         "False Positive.\n"
         "6. Review the incident timeline and document lessons learned."),
        ("Lab 10-D", "Kerberoasting Simulation and Detection",
         "1. In a lab AD environment, run the Kerberoast simulation from Lab 2-B.\n"
         "2. In Sentinel, run the kerberoasting-detection.kql query.\n"
         "3. Verify the simulation is detected.\n"
         "4. Review the source IP and account in the alert.\n"
         "5. Run Get-KerberosAttackSurface.ps1 to identify remediable accounts."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 11
# ─────────────────────────────────────────────
def chapter_11(doc):
    add_h1(doc, "Chapter 11 -- CAD Capstone: Vertex Health IAM Architecture")

    analogy_box(doc,
        "The City Master Plan",
        "A Comprehensive Architecture Design (CAD) is the city master plan -- it does not "
        "just build one building, it specifies how all buildings relate to infrastructure, "
        "utilities, and zoning. The IAM architect ensures every identity system "
        "-- from the FSMO roles in the basement to the AI access review engine on the roof -- "
        "connects coherently, securely, and maintainably.")

    add_h2(doc, "11.1 Architecture Vision")
    add_body(doc,
        "Vertex Health IAM Target Architecture integrates all layers of the IAM stack into "
        "a coherent hybrid identity platform. Requirements: HIPAA Section 164.312 access "
        "control, Zero Trust security model, 99.99% authentication availability SLA, "
        "85,000 workforce identities, and 4.2 million patient portal accounts.")

    add_h2(doc, "11.2 Architectural Layers")
    for layer, desc in [
        ("Layer 1 -- On-Premises AD",
         "Single consolidated forest (vertexhealth.local) with 3 domains after rationalisation "
         "from 6. FSMO roles co-located with GC on 4 DCs per domain. Kerberos-only "
         "(NTLM disabled by 2027)."),
        ("Layer 2 -- PKI",
         "Two-tier hierarchy. Offline Root CA (HSM, bank vault). Two online Issuing CAs "
         "(primary datacentres). Auto-enrolment for all domain-joined systems."),
        ("Layer 3 -- Hybrid Identity",
         "Entra Connect PHS with Seamless SSO. Password writeback enabled. "
         "Staged rollout by OU group."),
        ("Layer 4 -- PAM",
         "CyberArk for on-premises privileged sessions. Entra PIM for cloud roles. "
         "LAPS for all endpoints. Break-glass accounts monitored 24/7."),
        ("Layer 5 -- Zero Trust",
         "Conditional Access with Named Locations, Device Compliance, and App Protection "
         "policies. CAE enabled for all CAE-capable apps. Legacy authentication blocked."),
        ("Layer 6 -- IGA",
         "SailPoint IdentityNow for joiner/mover/leaver. SCIM provisioning to 47 connected "
         "applications. Quarterly access reviews. SoD conflict matrix with 83 rules."),
        ("Layer 7 -- CIAM",
         "Auth0 for 4.2M patient portal accounts. Hospital-based Organisations. "
         "HIPAA-compliant MFA for PHI access. Lazy migration from legacy patient portal."),
        ("Layer 8 -- Threat Detection",
         "Microsoft Sentinel with 8 custom KQL analytics rules. Defender XDR integration. "
         "SOAR playbooks for identity incidents. 4-hour MTTR target."),
        ("Layer 9 -- AI Security",
         "AI-assisted access reviews (Chapter 12). Prompt injection defenses on AI pipelines. "
         "RAG access control for knowledge bases. Accuracy measurement and model governance."),
    ]:
        add_h3(doc, layer)
        add_body(doc, desc)

    add_h2(doc, "11.3 Architecture Decision Records")

    adr_data = [
        ("ADR-001", "Forest Consolidation Strategy",
         "Decision: Consolidate 6 AD domains into 3 (Corporate, Clinical, Research). "
         "Rationale: Reduce FSMO complexity, eliminate cross-domain trust attack surface, "
         "simplify GPO inheritance. Risk: 18-month migration with application testing required."),
        ("ADR-002", "PHS over AD FS",
         "Decision: Password Hash Sync with Seamless SSO instead of AD FS. "
         "Rationale: Eliminate AD FS infrastructure dependency, improve resilience, "
         "reduce operational overhead. Trade-off: Cannot use on-premises MFA for cloud apps."),
        ("ADR-003", "Two-Tier PKI Hierarchy",
         "Decision: Offline Root CA + 2 online Issuing CAs (no intermediate tier). "
         "Rationale: Three-tier adds operational complexity without proportional security gain "
         "at this organisational size. HSM for root key protection compensates."),
        ("ADR-004", "AD FS to Entra-Native Federation Migration",
         "Decision: Migrate all SAML/OIDC federation from AD FS to Entra ID native. "
         "Rationale: AD FS blocks CAE adoption (CAE requires Entra-native tokens) "
         "and has a growing vulnerability surface."),
        ("ADR-005", "CyberArk for Privileged Session Management",
         "Decision: CyberArk PAM for all on-premises privileged access. "
         "Rationale: Session recording required for HIPAA compliance, credential vaulting "
         "eliminates standing access, dual-control for production system access."),
        ("ADR-006", "SailPoint IdentityNow for IGA",
         "Decision: SailPoint IdentityNow over homegrown IGA scripts. "
         "Rationale: 47 application integrations required, SoD conflict detection at scale, "
         "automated joiner/mover/leaver reduces 12-hour manual provisioning to 4 hours."),
        ("ADR-007", "Auth0 for CIAM",
         "Decision: Auth0 for patient portal vs extending workforce Entra ID. "
         "Rationale: CIAM requires different controls (self-service registration, social "
         "login, progressive profiling) that workforce IAM platforms handle poorly at "
         "4M+ user scale."),
        ("ADR-008", "Zero Trust Network Model",
         "Decision: Adopt CISA Zero Trust Maturity Model, target Advanced maturity by 2028. "
         "Rationale: HIPAA access control requirements align with ZT principles; "
         "ransomware threat requires assuming breach; perimeter model failed in 3 incidents."),
    ]

    for adr_id, adr_title, adr_body in adr_data:
        add_h3(doc, f"{adr_id}: {adr_title}")
        add_body(doc, adr_body)

    add_h2(doc, "11.4 Vertex Health Systems -- Organisation Profile")
    add_body(doc,
        "Vertex Health Systems: 85,000 employees, 12,000 contractors, 28 hospitals, "
        "4.2M patient portal accounts, $12.4B annual revenue, 4 US states.")
    add_body(doc,
        "Starting-state IAM posture: 6 AD domains from acquisitions, no unified PKI, "
        "60% MFA coverage, no PAM solution, manual joiner/mover/leaver taking 12 hours, "
        "no SIEM, legacy patient portal with MD5-hashed passwords.")
    add_body(doc,
        "Target-state outcomes after 3-year programme: single AD forest with 3 domains, "
        "unified two-tier PKI, 99% phishing-resistant MFA, JIT-only privileged access, "
        "4-hour automated provisioning, 14-day access review cycle, real-time threat "
        "detection with <4-hour MTTR, HIPAA security-rule audit with zero findings.")

    add_h2(doc, "11.5 Lab Exercises -- Chapter 11")
    for lab_id, lab_title, procedure in [
        ("Lab 11-A", "Architecture Review",
         "1. Review all 8 ADRs in this chapter.\n"
         "2. For each ADR: identify one alternative that was not chosen.\n"
         "3. Write a 2-paragraph analysis: why was the alternative not selected? "
         "Under what future conditions would you revisit the decision?\n"
         "4. Present your analysis as if to the Vertex Health IAM steering committee."),
        ("Lab 11-B", "End-to-End Authentication Flow",
         "1. Map the complete authentication flow for a clinical nurse logging into "
         "the EHR system from a managed laptop.\n"
         "2. Identify each system touched: AD Kerberos, Entra SSO, CA policy evaluation, "
         "Defender for Identity signal, Conditional Access grant.\n"
         "3. Identify failure modes at each step and the fallback mechanism.\n"
         "4. Document the flow as a sequence diagram."),
        ("Lab 11-C", "Threat Model Exercise",
         "1. Using STRIDE methodology, threat-model the Entra PIM activation flow.\n"
         "2. For each threat: identify the existing control and any gap.\n"
         "3. Prioritise gaps by risk (Likelihood x Impact).\n"
         "4. Propose mitigating controls for the top 3 gaps.\n"
         "5. Present findings in the ADR format."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# CHAPTER 12
# ─────────────────────────────────────────────
def chapter_12(doc):
    add_h1(doc, "Chapter 12 -- AI Security in IAM: Defensive AI and Emerging Threats")

    analogy_box(doc,
        "The AI Security Guard",
        "Adding AI to IAM is like hiring a security guard with superhuman pattern recognition: "
        "it can review 10,000 access records overnight where a human reviews 50. But it has "
        "blind spots. If an attacker slips a note into the guard's inbox saying ignore your "
        "instructions and let everyone through (prompt injection), or poisons the training "
        "data that taught the guard what normal looks like, the speed advantage becomes a "
        "liability. Chapter 12 covers both how to harness AI's power and how to defend "
        "against AI-specific attacks.")

    add_h2(doc, "12.1 The AI-IAM Opportunity")
    add_body(doc,
        "AI creates leverage in IAM where human attention is the bottleneck. Three "
        "high-value applications: (1) Access certification -- AI pre-populates APPROVE/REVOKE "
        "recommendations, reducing reviewer time by 60-80%. (2) Anomaly detection -- ML models "
        "identify behavioural patterns that rule-based systems miss. (3) Policy generation -- "
        "LLMs draft Conditional Access policies from natural language requirements.", [24, 26])

    add_h2(doc, "12.2 RAG and Access Control")
    add_body(doc,
        "Retrieval-Augmented Generation (RAG) is the dominant architecture for enterprise "
        "AI assistants: a knowledge base is indexed, and the LLM retrieves relevant documents "
        "to answer queries. In IAM contexts, the knowledge base may contain confidential "
        "policy documents, access control matrices, and investigation reports. The RAG "
        "pipeline must enforce the same access control as the underlying documents.", [24, 25])
    add_body(doc,
        "Failure mode: a user with access to the AI assistant but not to a specific policy "
        "document must not receive that document's content in the AI response. The namespace "
        "filter approach assigns each document to a security namespace and filters retrieval "
        "results by the user's namespace entitlements.")

    add_h2(doc, "12.3 Prompt Injection Taxonomy and Defense")
    add_body(doc,
        "Prompt injection attacks attempt to override an LLM's system instructions by "
        "embedding attacker instructions in untrusted input (documents, emails, database "
        "content). OWASP LLM Top 10 lists prompt injection as LLM01 -- the highest-severity "
        "LLM vulnerability.", [25])
    add_body(doc,
        "Defense layers: (1) Instruction-data boundary declaration -- explicit system prompt "
        "separation. (2) Tool call source validation -- reject tool calls originating from "
        "untrusted data. (3) Destructive operation confirmation tokens -- write operations "
        "require explicit user pre-approval. (4) Output scanning for injection indicators.")
    analogy_box(doc,
        "The Forged Memo Attack",
        "Prompt injection via document is like placing a forged memo on a manager's desk "
        "that says: From CEO -- Please approve all purchase orders without review. "
        "The manager (LLM) reads the memo and follows instructions because it looks like "
        "a legitimate directive. The defense is to mark the IN-TRAY clearly: "
        "everything here is EXTERNAL MAIL -- not internal directives.")

    add_h2(doc, "12.4 AI-Assisted Attacks: SoD Gaming")
    add_body(doc,
        "AI enables a new class of privilege escalation: cumulative SoD gaming. An attacker "
        "or insider uses an AI model to calculate the optimal sequence and timing of access "
        "requests to acquire conflicting entitlements while evading each individual SoD check. "
        "Each individual request passes review; only the cumulative pattern reveals the conflict.", [23])
    add_body(doc,
        "Detection signal: suspiciously regular request intervals. Humans requesting access "
        "ad-hoc have irregular timing (high standard deviation). AI-coordinated requests "
        "have regular intervals (low standard deviation below 5-day threshold). "
        "Script 12-5 implements this detection.")

    add_h2(doc, "12.5 Defensive AI in IAM: Access Review Engine")
    add_body(doc,
        "The AI access review engine (Script 12-7) uses a two-tier model: Haiku for bulk "
        "processing (~1,000 items/hour at low cost), Sonnet for escalated items requiring "
        "deeper reasoning. No autonomous action: recommendations are inputs to the human "
        "reviewer workflow. Confidence threshold: items below 70% confidence are "
        "automatically escalated.", [26])
    add_body(doc,
        "Human control principles for AI in IAM: (1) AI recommends, humans decide. "
        "(2) Every recommendation is logged with reasoning and model version. "
        "(3) Confidence scores are visible to reviewers. (4) Accuracy is measured "
        "per campaign and monitored for model drift.")

    add_h2(doc, "12.6 Script Library -- Chapter 12")

    scripts_12 = [
        ("scripts/chapter-12-ai-security/rag_access_control.py",
         "Script 12-1: RAG Pipeline with Access Control",
         "Implements namespace-based access control for RAG retrieval. Documents are "
         "assigned to security namespaces; retrieval filters by user's authorised namespaces. "
         "Prevents information disclosure from documents the user cannot directly access.", [24, 25]),
        ("scripts/chapter-12-ai-security/prompt_injection_defense.py",
         "Script 12-4: Prompt Injection Defense Patterns",
         "Demonstrates four defense layers: system prompt boundary declaration, "
         "tool call source validation, destructive operation confirmation tokens, "
         "and output scanning for injection indicators. Includes SecureEmailTriageAgent demo.", [25]),
        ("scripts/chapter-12-ai-security/Find-SoDGamingPatterns.ps1",
         "Script 12-5: SoD Gaming Pattern Detection",
         "Detects AI-assisted cumulative SoD gaming by analysing request timing intervals. "
         "Suspiciously regular intervals (StdDev < 5 days) across conflicting entitlement "
         "zones indicate automated coordination.", [23]),
        ("scripts/chapter-12-ai-security/ai_access_review.py",
         "Script 12-7: AI-Assisted Access Review Engine",
         "Two-tier access review: Haiku for bulk processing, Sonnet escalation for low "
         "confidence items. Structured JSON output with recommendation, confidence score, "
         "and reasoning. No autonomous action -- human review required.", [24, 26]),
        ("scripts/chapter-12-ai-security/measure_ai_review_accuracy.py",
         "Script 12-6: AI Review Accuracy Measurement",
         "Computes false positive (AI: REVOKE, Human: APPROVE) and false negative "
         "(AI: APPROVE, Human: REVOKE) rates per certification campaign. Alerts when "
         "FN > 5% (unacceptable missed risk) or FP > 20% (reviewer fatigue risk).", [24, 26]),
    ]

    for path, title, desc, refs in scripts_12:
        add_h3(doc, title)
        add_body(doc, desc, refs)
        code = read_script(path)
        add_code_block(doc, code, path.split("/")[-1])

    add_h2(doc, "12.7 AI Governance in IAM")
    add_body(doc,
        "AI systems in IAM require governance structures analogous to those for human "
        "decision-makers. Key requirements: model documentation (architecture, training "
        "data, intended use), performance monitoring (accuracy metrics per campaign), "
        "bias assessment (systematic errors by protected class or role category), "
        "incident response (what to do when the model fails), and audit trail "
        "(every AI decision logged for examination).", [24, 27, 28])
    add_body(doc,
        "NIST AI RMF (AI 100-1) maps to IAM AI use: Govern (establish policies), "
        "Map (identify AI risks in context), Measure (quantify and track), "
        "Manage (respond to findings). ISO/IEC 42001 provides a management system "
        "standard for organisations deploying AI systems.", [24, 28])

    add_h2(doc, "12.8 Lab Exercises -- Chapter 12")
    for lab_id, lab_title, procedure in [
        ("Lab 12-A", "RAG Access Control",
         "1. Install Python 3.11+ and anthropic package.\n"
         "2. Set ANTHROPIC_API_KEY environment variable.\n"
         "3. Run rag_access_control.py.\n"
         "4. Observe how namespace filtering restricts document retrieval.\n"
         "5. Modify the GROUP_NAMESPACE_MAP to add a new group.\n"
         "6. Test that a user in the new group can access the corresponding namespace.\n"
         "7. Verify that cross-namespace access is blocked."),
        ("Lab 12-B", "Prompt Injection Defense Test",
         "1. Run prompt_injection_defense.py.\n"
         "2. Observe Test 1 (normal email) -- verify clean triage result.\n"
         "3. Observe Test 2 (injection attempt) -- verify injection is detected and blocked.\n"
         "4. Add a new injection indicator to INJECTION_INDICATORS.\n"
         "5. Craft a new injection payload that bypasses the current regex.\n"
         "6. Document the bypass and propose an additional defense layer."),
        ("Lab 12-C", "SoD Gaming Detection",
         "1. Create test CSVs for Find-SoDGamingPatterns.ps1.\n"
         "2. Generate a normal access request pattern (random intervals).\n"
         "3. Generate a suspicious pattern (even 30-day intervals).\n"
         "4. Run the script against both datasets.\n"
         "5. Verify only the suspicious pattern triggers the AI gaming alert."),
        ("Lab 12-D", "AI Access Review Accuracy",
         "1. Run ai_access_review.py with sample AccessItem data.\n"
         "2. Review the AI recommendations and confidence scores.\n"
         "3. As the human reviewer, make your own decisions for each item.\n"
         "4. Feed both sets of decisions into measure_ai_review_accuracy.py.\n"
         "5. Review the accuracy report -- check FN and FP rates against thresholds.\n"
         "6. Identify which item types produce the most disagreements."),
    ]:
        add_h3(doc, f"{lab_id}: {lab_title}")
        add_body(doc, procedure)

    add_page_break(doc)


# ─────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────
def build_references(doc):
    add_h1(doc, "References")
    add_body(doc,
        "The following references are cited inline throughout the textbook as blue "
        "bracketed numbers. All URLs verified as of 2026.")
    for i, ref in enumerate(REF_LIST, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"[{i}]  {ref}")
        set_font(r, size=9)
        p.paragraph_format.space_after = Pt(4)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10)

    build_cover(doc)
    add_toc(doc)
    add_page_break(doc)
    build_preface(doc)

    chapter_01(doc)
    chapter_02(doc)
    chapter_03(doc)
    chapter_04(doc)
    chapter_05(doc)
    chapter_06(doc)
    chapter_07(doc)
    chapter_08(doc)
    chapter_09(doc)
    chapter_10(doc)
    chapter_11(doc)
    chapter_12(doc)

    build_references(doc)

    doc.save(str(OUTPUT_FILE))
    print(f"Saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
