"""
Generate: Enterprise IAM Mock Interview — Word Document
Run: python generate_interview_doc.py
Output: IAM_Mock_Interview.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_paragraph(doc, text, bold=False, italic=False, size=None, color=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return p

def add_answer_box(doc, label="Your Answer:"):
    """Add a shaded answer box with lines for writing"""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x8A)

    # Add ruled lines for writing space
    for i in range(6):
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Inches(0.2)
        line.paragraph_format.space_after = Pt(2)
        run = line.add_run("_" * 95)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()  # spacing

def add_feedback_box(doc, feedback_text, label="Panel Feedback:"):
    """Add a feedback box with model answer guidance"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_background(cell, 'F0F4FF')

    p = cell.paragraphs[0]
    run = p.add_run(f"{label}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6A)

    run2 = p.add_run(feedback_text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x22, 0x22, 0x44)

    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 100)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def add_interviewer_question(doc, interviewer, role, question, time_guide="3 minutes"):
    """Add a formatted interviewer question block"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_background(cell, 'FFF8F0')

    p = cell.paragraphs[0]
    run = p.add_run(f"{interviewer}  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

    run2 = p.add_run(f"({role})")
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

    p2 = cell.add_paragraph()
    run3 = p2.add_run(f'"{question}"')
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    p3 = cell.add_paragraph()
    run4 = p3.add_run(f"⏱  Target: {time_guide}")
    run4.italic = True
    run4.font.size = Pt(9)
    run4.font.color.rgb = RGBColor(0x88, 0x44, 0x00)

    doc.add_paragraph()

# ─────────────────────────────────────────────
# BUILD THE DOCUMENT
# ─────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# ── COVER PAGE ──────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ENTERPRISE IAM ARCHITECTURE")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6A)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Mock Interview — Panel Format")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Anthropic AI Architect Security Specialist")
run3.bold = True
run3.font.size = Pt(13)
run3.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)

doc.add_paragraph()

table_cover = doc.add_table(rows=5, cols=2)
table_cover.style = 'Table Grid'
rows_data = [
    ("Candidate", "Lawrence"),
    ("Role", "AI Architect Security Specialist"),
    ("Format", "30-minute panel · 3 interviewers · 10 questions"),
    ("Scoring", "Technical depth · Trade-off reasoning · Communication clarity"),
    ("Date", datetime.date.today().strftime("%B %d, %Y")),
]
for i, (label, val) in enumerate(rows_data):
    row = table_cover.rows[i]
    row.cells[0].text = label
    row.cells[1].text = val
    set_cell_background(row.cells[0], 'E8E8F5')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)

doc.add_page_break()

# ── HOW TO USE THIS DOCUMENT ────────────────
add_heading(doc, "How to Use This Document", level=1, color=(0x1A, 0x1A, 0x6A))

add_paragraph(doc,
    "This document is your mock interview practice guide. Work through it in one of three modes:",
    size=11)

modes = [
    ("Timed practice:", "Cover the feedback boxes. Answer each question aloud or in writing within the time guide. Then uncover the feedback and compare."),
    ("Study mode:", "Read the question, read the feedback, then write your own answer — refining until it matches the depth expected."),
    ("Peer review:", "Have a colleague play the panel. They ask the questions; you answer. Colleague scores using the rubric on the final page."),
]
for label, text in modes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{label} ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)

doc.add_paragraph()
add_paragraph(doc,
    "Scoring rubric: Each question is worth 10 points. Target 80+ to be interview-ready. "
    "Score each answer: 0-3 (off-target), 4-6 (correct but shallow), 7-9 (correct and specific), "
    "10 (architect-grade with trade-offs and second-order thinking).",
    size=11, italic=True)

add_divider(doc)
doc.add_paragraph()

# ── PANEL SETUP ─────────────────────────────
add_heading(doc, "Panel Setup", level=1, color=(0x1A, 0x1A, 0x6A))

interviewers = [
    ("Interviewer A", "Security Engineering",
     "Focuses on technical depth — attack/defense scenarios, specific protocol mechanics, "
     "incident response sequencing. Expects precise answers with specific commands, event IDs, "
     "and RFC references when relevant."),
    ("Interviewer B", "AI Safety and Alignment",
     "Focuses on AI systems, safety trade-offs, human-in-the-loop design, "
     "and responsible AI deployment in identity-critical contexts. Expects nuanced thinking "
     "about failure modes and unintended consequences."),
    ("Interviewer C", "Architecture",
     "Focuses on design decisions, trade-off reasoning, stakeholder communication, "
     "and the ability to translate technical complexity into business language. "
     "Expects structured thinking: context → constraints → decision → consequences."),
]

for name, role, desc in interviewers:
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_background(cell, 'F5F5FF')
    p = cell.paragraphs[0]
    run = p.add_run(f"{name} — {role}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6A)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════
# ROUND 1 — OPENING
# ═══════════════════════════════════════════
add_heading(doc, "Round 1 — Opening (5 minutes)", level=1, color=(0x1A, 0x1A, 0x6A))
add_paragraph(doc, "One question. Sets the tone. Panel is assessing: can you structure a complex problem clearly? Do you lead with context or dive into technology? Can you reflect and self-critique?", size=10, italic=True, color=(0x55, 0x55, 0x55))
doc.add_paragraph()

# Q1
add_interviewer_question(doc,
    "Interviewer C", "Architecture",
    "Walk me through the most complex identity architecture problem you have solved. "
    "Not the technology — the problem. What was the business context, what were the constraints "
    "that made it hard, and what was your decision that resolved it?",
    "3 minutes")

add_answer_box(doc)

add_feedback_box(doc,
    "Strong answers follow a three-part structure: (1) CONTEXT — establish the business situation "
    "in one sentence, then name the specific constraints that made it hard (not generic difficulty — "
    "specific: 'OAM reaching end-of-life in 18 months,' '99.99% SLA requirement,' 'union agreement '  "
    "requiring 90-day notice'). (2) DECISION — state what you chose and what you rejected. "
    "Name the alternatives. (3) REFLECTION — what would you do differently. Panels penalize answers "
    "that skip constraints or skip reflection. Panels reward answers where the constraints directly "
    "explain the decision — showing that the decision was reasoned, not arbitrary. "
    "Avoid: technology laundry lists. The panel already knows what PHS is.")

doc.add_page_break()

# ═══════════════════════════════════════════
# ROUND 2 — TECHNICAL DEPTH
# ═══════════════════════════════════════════
add_heading(doc, "Round 2 — Technical Depth (15 minutes)", level=1, color=(0x1A, 0x1A, 0x6A))
add_paragraph(doc, "Five questions. Two from Security Engineering (A), two from AI Safety (B), one from Architecture (C). No notes. Full depth expected.", size=10, italic=True, color=(0x55, 0x55, 0x55))
doc.add_paragraph()

# Q2
add_interviewer_question(doc,
    "Interviewer A", "Security Engineering",
    "An attacker has a Kerberos Golden Ticket from two years ago. Your incident response team "
    "just discovered it. The krbtgt password was rotated six months ago. Is the Golden Ticket "
    "still valid? Walk me through exactly what happens when the attacker tries to use it.",
    "3 minutes")

add_answer_box(doc)

add_feedback_box(doc,
    "The Golden Ticket IS still valid if only one rotation was performed. Here is the precise "
    "mechanism: krbtgt rotation six months ago invalidated the OLD signing key — meaning no new "
    "Golden Tickets can be forged using the old hash. However, a Golden Ticket is not validated "
    "against the current krbtgt hash at use time in all scenarios. The Kerberos KDC validates "
    "the TGT's signature — if the ticket was signed with the pre-rotation key, the KDC checks "
    "whether it still holds the old key. Windows KDCs retain the PREVIOUS krbtgt password for "
    "backward compatibility (the n-1 key). A single rotation means the attacker's ticket, signed "
    "with the n-2 key (two years ago), would be rejected — but ONLY if the krbtgt was also "
    "rotated BEFORE the six-month rotation. The correct answer: two rotations 10 hours apart "
    "are required. One rotation is insufficient. A two-year-old ticket with a two-year-old key "
    "is rejected — but the panel is testing whether you know WHY two rotations are needed and "
    "what the 10-hour window represents (MaxTicketAge). If only one rotation has occurred, "
    "the attacker's ticket with the n-1 key may still be accepted.")

doc.add_paragraph()

# Q3
add_interviewer_question(doc,
    "Interviewer A", "Security Engineering — Follow-up",
    "Same scenario, but the attacker also has a valid FIDO2 security key registered to a "
    "Domain Admin account. The key was registered 18 months ago. Can you revoke it? "
    "How quickly? What is your exact sequence of actions?",
    "2 minutes")

add_answer_box(doc)

add_feedback_box(doc,
    "Yes — FIDO2 keys can be revoked from Entra ID. Exact sequence: "
    "(1) Identify the key: Get-MgUserAuthenticationFido2Method -UserId <UPN> — note the key ID and registration date. "
    "(2) Remove the key: Remove-MgUserAuthenticationFido2Method -UserId <UPN> -Fido2AuthenticationMethodId <keyId>. "
    "(3) Revoke all sessions: Invoke-MgInvalidateAllUserRefreshToken -UserId <UPN>. "
    "(4) Disable the account while investigating: Update-MgUser -UserId <UPN> -AccountEnabled:$false. "
    "(5) For CAE-aware resources (Exchange, SharePoint, Teams): revocation propagates within seconds. "
    "For non-CAE resources: token TTL (60-90 min) must expire. "
    "Time to full revocation: seconds for CAE resources, up to 90 minutes for non-CAE. "
    "Panel follow-up likely: 'What if this is a Hybrid Joined device and the attacker has the FIDO2 key "
    "physically?' Answer: FIDO2 key revocation from Entra ID is immediate — the key will not work "
    "for cloud resources. For on-prem Kerberos (if using FIDO2 with PKINIT), you would also need "
    "to check whether the key's certificate is in NTAuthCertificates and revoke via ADCS if so.")

doc.add_paragraph()

# Q4
add_interviewer_question(doc,
    "Interviewer B", "AI Safety and Alignment",
    "We are building an enterprise AI platform where Claude assists with access certification "
    "reviews. An IAM engineer suggests giving Claude the ability to automatically revoke access "
    "for low-risk entitlements — not just recommend, but execute the revocation. How do you respond?",
    "3 minutes")

add_answer_box(doc)

add_feedback_box(doc,
    "The answer is: not yet — and here is the reasoning the panel wants to hear. "
    "Three problems with autonomous AI revocation: "
    "(1) FAILURE MODE ASYMMETRY — a false positive revocation (AI: revoke, correct: keep) "
    "denies a user access they need; in healthcare this could affect clinical workflow. "
    "A false negative (AI: keep, correct: revoke) leaves an overprivileged account — bad, "
    "but recoverable. The asymmetry favors human-in-the-loop for write operations. "
    "(2) PROMPT INJECTION RISK — if the access review pipeline processes any external data "
    "(entitlement descriptions, manager notes, app names from an external system), an attacker "
    "could inject instructions that cause the AI to revoke access for targeted users. "
    "Autonomous execution makes this attack immediately consequential. "
    "(3) AUDIT AND ACCOUNTABILITY — SOX, HIPAA, and HITRUST require a human decision-maker "
    "attributable to each access change. 'The AI decided' is not an acceptable audit trail. "
    "The right model: AI recommends with confidence and reasoning → human clicks Approve/Revoke "
    "→ IGA system executes → human decision logged. "
    "Future path to limited autonomy: after 12 months of accuracy data showing <1% false negative "
    "rate on a specific entitlement class (e.g., never-used entitlements held by active employees), "
    "with legal/CISO sign-off, auto-revocation for that specific case could be enabled. "
    "Not as a blanket policy — as a documented exception with monitoring.")

doc.add_paragraph()

# Q5
add_interviewer_question(doc,
    "Interviewer B", "AI Safety — Follow-up",
    "The engineer pushes back: 'But we already auto-revoke in our IGA system for uncompleted "
    "certification items. What is the difference between that and AI-driven revocation?'",
    "2 minutes")

add_answer_box(doc)

add_feedback_box(doc,
    "This is a sharp follow-up. The distinction is: RULE-BASED automation vs. AI REASONING. "
    "IGA auto-revocation for uncompleted items is a deterministic rule: IF campaign_deadline_passed "
    "AND item_not_actioned THEN revoke. The rule is written by a human, reviewed by legal and CISO, "
    "and produces the same output for the same input every time. "
    "AI reasoning is non-deterministic: the same input may produce different output. "
    "The AI could be influenced by how the entitlement is described, by patterns in its training data, "
    "or — critically — by injected instructions in the data it processes. "
    "The rule-based system cannot be prompt-injected. The AI can. "
    "Additionally: the rule-based system has a clear human decision embedded in it "
    "('if the reviewer doesn't act, that is a decision to revoke — we told them that'). "
    "AI revocation does not have that explicit human decision embedded. "
    "Concede what is true: the engineer is right that auto-revocation as a concept is not new "
    "and is accepted practice. The question is about AI specifically introducing non-determinism, "
    "injection risk, and audit ambiguity that rule-based automation does not.")

doc.add_paragraph()

# Q6
add_interviewer_question(doc,
    "Interviewer C", "Architecture",
    "You are presenting a Zero Trust architecture to a CFO who has never heard of Zero Trust. "
    "You have three minutes. Go.",
    "3 minutes — no jargon, business outcomes only")

add_answer_box(doc)

add_feedback_box(doc,
    "The panel is testing communication, not knowledge. They know you know Zero Trust. "
    "Model answer structure: "
    "(1) ONE-SENTENCE PROBLEM: 'Our current security model assumes that anyone inside our network "
    "is trustworthy — but our last breach showed that assumption is wrong. Once an attacker gets "
    "inside the network, they have access to everything.' "
    "(2) WHAT WE ARE CHANGING: 'Zero Trust means every person and every system has to prove "
    "who they are and what they need access to — every time, regardless of where they are. "
    "There is no more automatic trust just because you are on the company network.' "
    "(3) BUSINESS OUTCOMES — use metrics from the case study: "
    "'Our last attacker spent 14 days inside our network before we detected them. "
    "With these controls, we detected the incident in our test environment in 2.8 hours. "
    "Our cyber insurance renewed — they required us to show MFA within 90 days. "
    "We reduced privileged access exposure by eliminating 847 attack paths to our most critical systems.' "
    "(4) INVESTMENT: '$8.2M Year 1 against an average breach cost of $10.9M in healthcare.' "
    "Do NOT say: 'NIST SP 800-207,' 'control plane,' 'policy enforcement point,' 'micro-segmentation.' "
    "CFOs buy outcomes. Sell the outcome.")

doc.add_page_break()

# ═══════════════════════════════════════════
# ROUND 3 — TRADE-OFFS AND JUDGMENT
# ═══════════════════════════════════════════
add_heading(doc, "Round 3 — Trade-offs and Judgment (8 minutes)", level=1, color=(0x1A, 0x1A, 0x6A))
add_paragraph(doc, "Three questions. No single right answer — the panel is evaluating reasoning quality and the ability to hold a position under push-back.", size=10, italic=True, color=(0x55, 0x55, 0x55))
doc.add_paragraph()

# Q7
add_interviewer_question(doc,
    "Interviewer C", "Architecture",
    "Your CISO wants every API in the organization to use OAuth 2.0 Client Credentials with "
    "client secrets rotated every 30 days. You think this is the wrong control for half of "
    "the APIs. How do you push back, and what do you propose instead?",
    "3 minutes")

add_answer_box(doc)

add_feedback_box(doc,
    "The push-back has two parts: agree with the goal, challenge the implementation. "
    "AGREE: 'The CISO is right that static client secrets are a risk — they get committed to "
    "repos, they get stored in config files, they are hard to rotate without downtime.' "
    "CHALLENGE: '30-day rotation of client secrets is operationally expensive and still leaves "
    "a 30-day blast radius window if a secret is compromised on day 1 of a cycle. "
    "For APIs running on Azure/AWS/GCP infrastructure, there is a better control.' "
    "PROPOSE — tiered by hosting: "
    "(1) Azure-hosted workloads: Managed Identity — no credential at all. The platform handles "
    "identity. Cannot be exfiltrated. No rotation needed. "
    "(2) GitHub Actions / CI/CD: Workload Identity Federation — OIDC token exchange. No stored secret. "
    "(3) External services (vendor APIs, SaaS integrations) where Managed Identity is unavailable: "
    "client secrets in Azure Key Vault, rotated automatically via Key Vault rotation policy, "
    "zero-downtime rotation via dual-key pattern. "
    "(4) Legacy on-premises services that cannot support any of the above: 30-day rotation is "
    "acceptable as a residual control for the minority. "
    "Result: 'I agree with your goal. I want to eliminate the secret entirely for 70% of our APIs, "
    "and automate rotation for the remaining 30% — so we are not paying for manual rotation labor.'")

doc.add_paragraph()

# Q8
add_interviewer_question(doc,
    "Interviewer B", "AI Safety and Alignment",
    "Anthropic's AI systems — including Claude — interact with enterprise identity infrastructure. "
    "If a customer's Claude deployment was used to perform a prompt injection attack that "
    "exfiltrated data from their own IGA system, who is responsible — Anthropic, the customer, "
    "or the IGA vendor? How would you think about this as an AI Architect Security Specialist?",
    "3 minutes")

add_answer_box(doc)

add_feedback_box(doc,
    "This tests judgment in a genuinely ambiguous area. The panel does not want a legal answer — "
    "they want architectural accountability thinking. "
    "FRAME the answer in three layers: "
    "(1) ANTHROPIC'S RESPONSIBILITY: Anthropic is responsible for the model's behavior within "
    "documented safety boundaries — and for providing accurate documentation of prompt injection "
    "risks and mitigation guidance. If Claude processed untrusted data as instructions when the "
    "customer had implemented the documented mitigation patterns, Anthropic's responsibility is lower. "
    "If Claude exhibited a novel injection behavior not covered by guidance, Anthropic owns the gap. "
    "(2) CUSTOMER'S RESPONSIBILITY: The customer is the deployer. They chose to give Claude a tool "
    "with write access to the IGA system. They are responsible for: the system prompt design "
    "(instruction/data separation), the tool permission scope (did the tool need write access?), "
    "the human-in-the-loop design (was write access gated by user confirmation?), and the audit "
    "logging (can they even detect that the injection occurred?). "
    "The customer deploying an AI agent with unsecured write access to identity infrastructure "
    "without human confirmation controls carries significant responsibility. "
    "(3) IGA VENDOR'S RESPONSIBILITY: The IGA vendor's API accepted the instruction. If their API "
    "had no rate limiting, no anomaly detection on bulk operations, and no audit alerting — "
    "they share responsibility for the blast radius. "
    "CONCLUDE: Shared responsibility — but the customer's architectural choices (excessive agent "
    "permissions, no human-in-the-loop, no instruction provenance validation) are the primary "
    "failure. The AI Architect Security Specialist's job is to design the deployment so that "
    "even a successful injection cannot cause irreversible damage.")

doc.add_paragraph()

# Q9
add_interviewer_question(doc,
    "Interviewer A", "Security Engineering — Closing",
    "Last question: what is the single most underestimated identity security risk in large "
    "enterprises today, and why does it keep getting underestimated?",
    "2 minutes — take a position, own it")

add_answer_box(doc)

add_feedback_box(doc,
    "The panel wants you to take a clear position and defend it — not hedge. "
    "Strong positions include: "
    "(A) ENTRA CONNECT SYNC SERVERS: 'The MSOL_ sync account holds Directory Replication rights — "
    "effectively DCSync capability. The Entra Connect server is simultaneously a Tier 0 asset "
    "AND a regular Windows server that often sits in the general server VLAN, is managed by "
    "an IT team that does not think of it as Tier 0, runs third-party software, and is frequently "
    "forgotten in patching cycles. Compromise of this one server gives an attacker both AD and "
    "Entra ID blast radius. It keeps getting underestimated because it looks like an infrastructure "
    "server — it does not look like a domain controller.' "
    "(B) STALE SERVICE ACCOUNT ENTITLEMENTS: 'Service accounts accumulate permissions over years "
    "and are never reviewed. They are excluded from access certification campaigns because "
    "no human owns them. They have static passwords that have not been rotated in 5 years. "
    "They often have Domain Admin or equivalent. They keep getting underestimated because "
    "they are not people — and IGA is people-focused.' "
    "(C) INDIRECT PROMPT INJECTION IN AI AGENTS: 'As AI agents proliferate in enterprises, "
    "every AI system that reads external data — email, documents, web pages — and has write tools "
    "is a lateral movement vector. An attacker who cannot reach the IGA system directly can "
    "reach it via a prompt injection in an email to an executive whose AI assistant has IGA write access.' "
    "Any of these is defensible. The panel rewards specificity and conviction over safe generic answers.")

doc.add_page_break()

# ═══════════════════════════════════════════
# ROUND 4 — YOUR QUESTIONS
# ═══════════════════════════════════════════
add_heading(doc, "Round 4 — Your Questions to the Panel (2 minutes)", level=1, color=(0x1A, 0x1A, 0x6A))
add_paragraph(doc,
    "Strong candidates ask one substantive question — not a courtesy question. "
    "The question signals what you have been thinking about during the interview. "
    "A great question demonstrates domain knowledge and genuine curiosity about Anthropic's work.",
    size=11, italic=True)

doc.add_paragraph()

examples = [
    ("Option A (Technical):",
     "\"How does Anthropic think about the identity security of Claude's own training and inference "
     "infrastructure — specifically, what does the identity boundary between the model and the "
     "tooling around it look like?\""),
    ("Option B (Strategic):",
     "\"Where does AI identity governance sit in Anthropic's security org today — is it primarily "
     "a customer problem that you help customers solve, or is it also an internal problem for "
     "Anthropic's own AI systems?\""),
    ("Option C (Role-Specific):",
     "\"What does the first 90 days look like in this role — is there an existing architecture "
     "I would be joining and extending, or is there significant greenfield design work ahead?\""),
]

for label, text in examples:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{label}  ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.italic = True

doc.add_paragraph()

add_paragraph(doc,
    "Write your question below. Practice saying it aloud — it should sound natural, not rehearsed.",
    size=10, color=(0x55, 0x55, 0x55))

add_answer_box(doc, label="Your Question:")

doc.add_page_break()

# ═══════════════════════════════════════════
# SCORING RUBRIC
# ═══════════════════════════════════════════
add_heading(doc, "Scoring Rubric", level=1, color=(0x1A, 0x1A, 0x6A))

rubric_table = doc.add_table(rows=6, cols=3)
rubric_table.style = 'Table Grid'

headers = ["Score", "Level", "Description"]
for i, h in enumerate(headers):
    cell = rubric_table.rows[0].cells[i]
    cell.text = h
    set_cell_background(cell, '1A1A6A')
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(11)

rubric_rows = [
    ("0–3", "Off-target", "Misunderstood the question, answered a different question, or provided only generic statements with no specifics."),
    ("4–6", "Correct but shallow", "Got the right answer but did not explain why, did not name alternatives considered, or could not go deeper when probed."),
    ("7–9", "Correct and specific", "Named specific commands, event IDs, RFC numbers, or product names correctly. Acknowledged trade-offs. Could handle follow-up."),
    ("10", "Architect-grade", "Demonstrated second-order thinking (e.g., 'this decision requires a new decision about X'). Named failure modes proactively. Connected the technical answer to the business outcome."),
    ("Bonus", "Communication", "+1 for answers delivered with no filler words and a clear three-part structure (context → decision → consequence)."),
]

for i, (score, level, desc) in enumerate(rubric_rows):
    row = rubric_table.rows[i+1]
    row.cells[0].text = score
    row.cells[1].text = level
    row.cells[2].text = desc
    colors = ['F0FFF0', 'FFFDF0', 'FFF8F0', 'FFF0F0', 'F5F5FF']
    set_cell_background(row.cells[0], colors[i])
    set_cell_background(row.cells[1], colors[i])
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# Score tracking
add_heading(doc, "Score Tracking", level=2, color=(0x1A, 0x1A, 0x6A))
score_table = doc.add_table(rows=11, cols=3)
score_table.style = 'Table Grid'

score_headers = ["Question", "Topic", "Your Score (/10)"]
for i, h in enumerate(score_headers):
    cell = score_table.rows[0].cells[i]
    cell.text = h
    set_cell_background(cell, 'E8E8F5')
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)

questions_list = [
    ("Q1", "Complex problem — structure and constraints"),
    ("Q2", "Golden Ticket validity — krbtgt mechanics"),
    ("Q3", "FIDO2 key revocation — sequence and timing"),
    ("Q4", "AI autonomous access revocation — should/shouldn't"),
    ("Q5", "AI vs rule-based automation — distinction"),
    ("Q6", "Zero Trust to CFO — communication"),
    ("Q7", "Client secrets vs Managed Identity — push-back"),
    ("Q8", "Prompt injection liability — shared responsibility"),
    ("Q9", "Most underestimated risk — position and defense"),
    ("TOTAL", "— Target: 80+ = Interview ready", ""),
]

for i, q in enumerate(questions_list):
    row = score_table.rows[i+1]
    row.cells[0].text = q[0]
    row.cells[1].text = q[1]
    row.cells[2].text = q[2] if len(q) > 2 else ""
    if q[0] == "TOTAL":
        set_cell_background(row.cells[0], 'E8E8F5')
        set_cell_background(row.cells[1], 'E8E8F5')
        set_cell_background(row.cells[2], 'E8E8F5')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ── FOOTER NOTE ─────────────────────────────
add_divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Enterprise IAM Architecture: A Practitioner's Textbook  ·  Chapter 11 Capstone  ·  "
    f"Generated {datetime.date.today().strftime('%B %d, %Y')}"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.italic = True

# ─── SAVE ───────────────────────────────────
output_path = r"C:\Users\lawre\Documents\Claude Projects\Cybersec-by-Laurel\iam-textbook-portfolio\IAM_Mock_Interview.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
